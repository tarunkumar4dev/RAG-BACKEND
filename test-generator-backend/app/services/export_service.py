"""
Export Service v14 — Multi-Template Support (Classic / Modern / Compact / Colorful)

v14 changes:
  - "colorful" template now uses a distinct INSTITUTE-PAPER layout
    (layout_style="institute_paper") matching a reference institute exam
    format: Institute name header, "CLASS X — SUBJECT" line, optional
    Topic line, Teacher / Subject + Max Marks / Date / Time meta row, a
    signature multi-color section rule, inline "[marks]" per question (no
    separate column), MCQ options rendered as (a)/(b) two-column pairs,
    plain "SECTION A (desc)" headings, and a "— All the Best —" footer.
  - generate_pdf() / generate_docx() gain new optional params:
    teacher_name, institute_name, duration, topic — only rendered by the
    institute_paper layout; other templates ignore them (still accepted,
    so callers never break).
  - modern / classic / compact keep the original card-based layout
    (layout_style="card_based") — completely unaffected.

v13 features retained:
  - TEMPLATE_PRESETS: 4 selectable visual templates for PDF + DOCX export
  - get_available_templates() metadata endpoint
  - Semantic colors (Accountancy emerald, Statistics blue) are NOT
    template-driven — they carry meaning, not brand styling.

v12 / v11 features retained:
  - Matrix bracket rendering, question tables (Statistics), Accountancy
    answer tables, CBSE section grouping, manual question images, LaTeX/
    Unicode cleanup, paper date support, card-style question boxes.
"""

import io
import re
import base64
import logging
from typing import List, Optional
from datetime import datetime
from urllib.request import urlopen, Request
from urllib.error import URLError, HTTPError

logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════════
# CBSE Section Definitions
# ═══════════════════════════════════════════════════════════════════════

CBSE_SECTIONS_META = {
    "A": {"title": "Section A", "subtitle": "(1 mark each — MCQ / Assertion-Reason)", "marks": 1, "instruction": "All questions are compulsory. Each carries 1 mark."},
    "B": {"title": "Section B", "subtitle": "(2 marks each — Very Short Answer)", "marks": 2, "instruction": "All questions are compulsory. Each carries 2 marks."},
    "C": {"title": "Section C", "subtitle": "(3 marks each — Short Answer)", "marks": 3, "instruction": "All questions are compulsory. Each carries 3 marks."},
    "D": {"title": "Section D", "subtitle": "(5 marks each — Long Answer)", "marks": 5, "instruction": "All questions are compulsory. Each carries 5 marks."},
    "E": {"title": "Section E", "subtitle": "(4 marks each — Case Study Based)", "marks": 4, "instruction": "All questions are compulsory. Each carries 4 marks. Answer all sub-parts."},
}

ACCOUNTANCY_SECTIONS_META = {
    "A_1m":  {"title": "Part A", "subtitle": "(1 mark each — MCQ / Assertion-Reason)", "marks": 1, "instruction": "Questions carry 1 mark each."},
    "A_3m":  {"title": "Part A", "subtitle": "(3 marks each)", "marks": 3, "instruction": "Questions carry 3 marks each."},
    "A_4m":  {"title": "Part A", "subtitle": "(4 marks each)", "marks": 4, "instruction": "Questions carry 4 marks each."},
    "A_6m":  {"title": "Part A", "subtitle": "(6 marks each)", "marks": 6, "instruction": "Questions carry 6 marks each."},
    "B1_1m": {"title": "Part B (Option I)", "subtitle": "Analysis of Financial Statements — (1 mark each)", "marks": 1, "instruction": "Questions carry 1 mark each."},
    "B1_3m": {"title": "Part B (Option I)", "subtitle": "Analysis of Financial Statements — (3 marks each)", "marks": 3, "instruction": "Questions carry 3 marks each."},
    "B1_4m": {"title": "Part B (Option I)", "subtitle": "Analysis of Financial Statements — (4 marks each)", "marks": 4, "instruction": "Questions carry 4 marks each."},
    "B1_6m": {"title": "Part B (Option I)", "subtitle": "Analysis of Financial Statements — (6 marks each)", "marks": 6, "instruction": "Questions carry 6 marks each."},
}

SECTION_ORDER = ["A", "B", "C", "D", "E"]
ACCOUNTANCY_SECTION_ORDER = ["A_1m", "A_3m", "A_4m", "A_6m", "B1_1m", "B1_3m", "B1_4m", "B1_6m"]


# ═══════════════════════════════════════════════════════════════════════
# Template Presets  (v14: each preset now carries a "layout_style")
# ═══════════════════════════════════════════════════════════════════════
#
# layout_style:
#   "card_based"       → modern / classic / compact — existing renderer
#   "institute_paper"  → colorful — institute exam-paper renderer (v14)
#
# card_style:
#   "card"   → rounded/bordered box around each question
#   "flat"   → no box, just spacing
#   "stripe" → bordered box with a colored left accent stripe per section

TEMPLATE_PRESETS = {
    "modern": {
        "label": "Modern",
        "description": "Clean sans-serif with card-style questions. Our signature look.",
        "layout_style": "card_based",
        "font_body": "Helvetica",
        "font_bold": "Helvetica-Bold",
        "docx_font": "Calibri",
        "primary": "#1a1a2e",
        "secondary": "#4a4a6a",
        "muted": "#6b7280",
        "light_muted": "#9ca3af",
        "border": "#e5e7eb",
        "card_bg": "#F9FAFB",
        "card_border": "#E5E7EB",
        "correct": "#047857",
        "table_header_bg": "#f3f4f6",
        "table_header_text": "#1f2937",
        "card_style": "card",
        "header_style": "formal",
        "section_colors": None,
        "spacing_scale": 1.0,
        "margins_cm": (1.0, 1.0, 1.5, 1.5),
    },
    "classic": {
        "label": "Classic",
        "description": "Traditional serif exam-paper look — formal, no boxes.",
        "layout_style": "card_based",
        "font_body": "Times-Roman",
        "font_bold": "Times-Bold",
        "docx_font": "Times New Roman",
        "primary": "#1a1a2e",
        "secondary": "#3f3f3f",
        "muted": "#595959",
        "light_muted": "#7a7a7a",
        "border": "#cfcfcf",
        "card_bg": "#FFFFFF",
        "card_border": "#cfcfcf",
        "correct": "#1e5631",
        "table_header_bg": "#f0f0f0",
        "table_header_text": "#1a1a1a",
        "card_style": "flat",
        "header_style": "formal",
        "section_colors": None,
        "spacing_scale": 1.0,
        "margins_cm": (1.2, 1.2, 1.8, 1.8),
    },
    "compact": {
        "label": "Compact",
        "description": "Dense layout, smaller fonts — fits more on fewer pages.",
        "layout_style": "card_based",
        "font_body": "Helvetica",
        "font_bold": "Helvetica-Bold",
        "docx_font": "Arial",
        "primary": "#1a1a2e",
        "secondary": "#4a4a6a",
        "muted": "#6b7280",
        "light_muted": "#9ca3af",
        "border": "#e5e7eb",
        "card_bg": "#FFFFFF",
        "card_border": "#e5e7eb",
        "correct": "#047857",
        "table_header_bg": "#f3f4f6",
        "table_header_text": "#1f2937",
        "card_style": "flat",
        "header_style": "minimal",
        "section_colors": None,
        "spacing_scale": 0.55,
        "margins_cm": (0.7, 0.7, 1.1, 1.1),
    },
    "colorful": {
        "label": "Colorful",
        "description": "Institute-style paper — inline marks, 2-column MCQ options, colorful section accents.",
        "layout_style": "institute_paper",
        "font_body": "Helvetica",
        "font_bold": "Helvetica-Bold",
        "docx_font": "Arial",
        "primary": "#1a1a2e",
        "secondary": "#4a4a6a",
        "muted": "#6b7280",
        "light_muted": "#9ca3af",
        "border": "#e5e7eb",
        "card_bg": "#FFFFFF",
        "card_border": "#e5e7eb",
        "correct": "#047857",
        "table_header_bg": "#f3f4f6",
        "table_header_text": "#1f2937",
        "card_style": "stripe",
        "header_style": "banner",
        "banner_bg": "#EEF2FF",
        "section_colors": {
            "A": "#2563eb",
            "B": "#7c3aed",
            "C": "#ea580c",
            "D": "#db2777",
            "E": "#0d9488",
            "F": "#65a30d",
        },
        "spacing_scale": 1.0,
        "margins_cm": (1.0, 1.0, 1.5, 1.5),
    },
}

DEFAULT_TEMPLATE = "modern"


def get_available_templates() -> list:
    """Metadata list for a frontend template-picker dropdown."""
    return [
        {"id": key, "label": val["label"], "description": val.get("description", "")}
        for key, val in TEMPLATE_PRESETS.items()
    ]


def _get_template(name: Optional[str]) -> dict:
    if not name:
        return TEMPLATE_PRESETS[DEFAULT_TEMPLATE]
    key = str(name).strip().lower()
    if key not in TEMPLATE_PRESETS:
        logger.warning(f"Unknown template '{key}', falling back to '{DEFAULT_TEMPLATE}'")
    return TEMPLATE_PRESETS.get(key, TEMPLATE_PRESETS[DEFAULT_TEMPLATE])


def _section_color(tpl: dict, section_label: Optional[str]) -> str:
    """Per-section accent color for 'colorful' template; falls back to primary."""
    sc = tpl.get("section_colors")
    if sc and section_label and section_label in sc:
        return sc[section_label]
    return tpl["primary"]


def _sc(val: float, tpl: dict) -> float:
    """Scale a PDF spacing value (points) by the template's spacing_scale."""
    return max(1, val * tpl.get("spacing_scale", 1.0))


def _hexnc(hexstr: str) -> str:
    """Hex color without '#', uppercased — for docx OXML shading/border fills."""
    return hexstr.lstrip('#').upper()


def _rgb(hexstr: str):
    """Hex string -> docx RGBColor."""
    from docx.shared import RGBColor
    h = hexstr.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


# ═══════════════════════════════════════════════════════════════════════
# Unicode sub/super scripts
# ═══════════════════════════════════════════════════════════════════════

UNICODE_SUBSCRIPTS = {
    '₀': '0', '₁': '1', '₂': '2', '₃': '3', '₄': '4',
    '₅': '5', '₆': '6', '₇': '7', '₈': '8', '₉': '9',
    '₊': '+', '₋': '-', '₌': '=',
    'ₐ': 'a', 'ₑ': 'e', 'ₒ': 'o', 'ₓ': 'x', 'ₙ': 'n',
}

UNICODE_SUPERSCRIPTS = {
    '⁰': '0', '¹': '1', '²': '2', '³': '3', '⁴': '4',
    '⁵': '5', '⁶': '6', '⁷': '7', '⁸': '8', '⁹': '9',
    '⁺': '+', '⁻': '-', '⁼': '=',
    'ⁿ': 'n', 'ⁱ': 'i',
}

CHEMICAL_PATTERN = re.compile(r'([A-Z][a-z]?)(\d+)')


def _fix_unicode_scripts(text: str, use_tags: bool = True) -> str:
    if not text:
        return text
    result = text
    if use_tags:
        sub_pattern = '([' + ''.join(re.escape(k) for k in UNICODE_SUBSCRIPTS.keys()) + ']+)'
        def sub_replacer(m):
            chars = m.group(1)
            converted = ''.join(UNICODE_SUBSCRIPTS.get(c, c) for c in chars)
            return f'<sub>{converted}</sub>'
        result = re.sub(sub_pattern, sub_replacer, result)

        sup_pattern = '([' + ''.join(re.escape(k) for k in UNICODE_SUPERSCRIPTS.keys()) + ']+)'
        def sup_replacer(m):
            chars = m.group(1)
            converted = ''.join(UNICODE_SUPERSCRIPTS.get(c, c) for c in chars)
            return f'<super>{converted}</super>'
        result = re.sub(sup_pattern, sup_replacer, result)
    else:
        for uni, plain in UNICODE_SUBSCRIPTS.items():
            result = result.replace(uni, plain)
        for uni, plain in UNICODE_SUPERSCRIPTS.items():
            result = result.replace(uni, plain)
    return result


def _fix_chemical_formulas(text: str, use_tags: bool = True) -> str:
    if not text:
        return text
    parts = re.split(r'(\$[^$]+\$)', text)
    result_parts = []
    for part in parts:
        if part.startswith('$') and part.endswith('$'):
            result_parts.append(part)
        else:
            if use_tags:
                fixed = CHEMICAL_PATTERN.sub(lambda m: f'{m.group(1)}<sub>{m.group(2)}</sub>', part)
            else:
                fixed = part
            result_parts.append(fixed)
    return ''.join(result_parts)


# ═══════════════════════════════════════════════════════════════════════
# Matrix Bracket Converter
# ═══════════════════════════════════════════════════════════════════════

def _convert_matrix_brackets(text: str, use_tags: bool = False) -> str:
    """Convert [[a,b],[c,d]] notation to proper matrix box format."""
    if not text:
        return text

    if '[[' in text:
        logger.warning(f"MATRIX_IN: {text[:150]!r}")

    bracket_pattern = re.compile(r'\[\[(.*?)\]\]', re.DOTALL)

    def build_matrix(inner_content: str) -> str:
        rows_raw = re.split(r'\],\s*\[', inner_content)
        if len(rows_raw) <= 1:
            return f"[[{inner_content}]]"

        matrix_rows = []
        max_cols = 0
        for row in rows_raw:
            values = [v.strip() for v in row.split(',') if v.strip()]
            matrix_rows.append(values)
            max_cols = max(max_cols, len(values))

        if not matrix_rows:
            return f"[[{inner_content}]]"

        if use_tags:
            lines = ['┌' + ' ' * (max_cols * 6) + '┐']
            for row in matrix_rows:
                padded = row + [''] * (max_cols - len(row))
                line = '│ ' + '  '.join(f'{v:>4}' for v in padded) + ' │'
                lines.append(line)
            lines.append('└' + ' ' * (max_cols * 6) + '┘')
            return '<br/>' + '<br/>'.join(lines) + '<br/>'
        else:
            lines = ['┌' + ' ' * (max_cols * 5) + '┐']
            for row in matrix_rows:
                padded = row + [''] * (max_cols - len(row))
                line = '│ ' + '  '.join(f'{v:>3}' for v in padded) + ' │'
                lines.append(line)
            lines.append('└' + ' ' * (max_cols * 5) + '┘')
            return '\n'.join(lines)

    def replace_matrix(match):
        inner = match.group(1)
        if not inner or ',' not in inner:
            return match.group(0)
        return build_matrix(inner)

    result = bracket_pattern.sub(replace_matrix, text)

    if '[[' in text:
        logger.warning(f"MATRIX_OUT: {result[:150]!r}")

    return result


# ═══════════════════════════════════════════════════════════════════════
# LaTeX → Clean Text
# ═══════════════════════════════════════════════════════════════════════

SYMBOL_MAP = {
    r'\times': '×', r'\div': '÷', r'\pm': '±', r'\mp': '∓', r'\cdot': '·',
    r'\leq': '≤', r'\geq': '≥', r'\neq': '≠', r'\approx': '≈',
    r'\equiv': '≡', r'\sim': '~', r'\propto': '∝',
    r'\infty': '∞', r'\therefore': '∴', r'\because': '∵',
    r'\cup': '∪', r'\cap': '∩', r'\subset': '⊂', r'\supset': '⊃',
    r'\subseteq': '⊆', r'\supseteq': '⊇', r'\in': '∈', r'\notin': '∉',
    r'\emptyset': '∅', r'\forall': '∀', r'\exists': '∃',
    r'\rightarrow': '→', r'\leftarrow': '←', r'\Rightarrow': '⇒',
    r'\Leftarrow': '⇐', r'\leftrightarrow': '↔', r'\to': '→',
    r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ', r'\delta': 'δ',
    r'\epsilon': 'ε', r'\zeta': 'ζ', r'\eta': 'η', r'\theta': 'θ',
    r'\iota': 'ι', r'\kappa': 'κ', r'\lambda': 'λ', r'\mu': 'μ',
    r'\nu': 'ν', r'\xi': 'ξ', r'\pi': 'π', r'\rho': 'ρ',
    r'\sigma': 'σ', r'\tau': 'τ', r'\phi': 'φ', r'\chi': 'χ',
    r'\psi': 'ψ', r'\omega': 'ω',
    r'\Gamma': 'Γ', r'\Delta': 'Δ', r'\Theta': 'Θ', r'\Lambda': 'Λ',
    r'\Sigma': 'Σ', r'\Phi': 'Φ', r'\Psi': 'Ψ', r'\Omega': 'Ω',
    r'\degree': '°', r'\circ': '°', r'\nabla': '∇',
    r'\partial': '∂', r'\ell': 'ℓ',
    r'\sum': 'Σ', r'\prod': 'Π', r'\int': '∫',
    r'\left': '', r'\right': '',
    r'\bigl': '', r'\bigr': '',
    r'\langle': '⟨', r'\rangle': '⟩',
    r'\lfloor': '⌊', r'\rfloor': '⌋', r'\lceil': '⌈', r'\rceil': '⌉',
    r'\setminus': ' \\ ',
}

TRIG_FUNCS = {
    r'\sin': 'sin', r'\cos': 'cos', r'\tan': 'tan',
    r'\cot': 'cot', r'\sec': 'sec', r'\csc': 'csc',
    r'\log': 'log', r'\ln': 'ln', r'\exp': 'exp',
    r'\lim': 'lim', r'\max': 'max', r'\min': 'min',
}

BARE_COMMANDS = {
    'setminus': ' \\ ', 'mathbb': '', 'mathrm': '', 'mathbf': '',
    'textbf': '', 'textit': '', 'overline': '', 'underline': '',
}

def _process_latex(text: str, use_tags: bool = False) -> str:
    if not text:
        return ""

    result = text
    result = result.replace('₹', 'Rs.')

    MODIFIER_LETTERS = {
        '\u1D57': 't', '\u02B0': 'h', '\u02E2': 's', '\u1D48': 'd',
        '\u02B3': 'r', '\u02E1': 'l', '\u1D43': 'a', '\u1D49': 'e',
        '\u1D52': 'o',
    }
    for mod, plain in MODIFIER_LETTERS.items():
        result = result.replace(mod, plain)

    result = _fix_chemical_formulas(result, use_tags)
    result = re.sub(r'\$([^$]+)\$', r'\1', result)

    mathbb_map = {'R': 'ℝ', 'Z': 'ℤ', 'N': 'ℕ', 'Q': 'ℚ', 'C': 'ℂ'}
    for letter, symbol in mathbb_map.items():
        result = result.replace(f'\\mathbb{{{letter}}}', symbol)
        result = result.replace(f'mathbb{{{letter}}}', symbol)
        result = re.sub(rf'(?<![a-zA-Z])mathbb\s*{letter}(?![a-zA-Z])', symbol, result)

    result = re.sub(r'\\(?:text|mathrm|mathbf|textbf|textit|mathit)\{([^}]*)\}', r'\1', result)
    result = re.sub(r'\\(?:overline|underline|bar|hat|tilde|vec)\{([^}]*)\}', r'\1', result)

    for _ in range(3):
        result = re.sub(r'\\frac\{([^{}]*)\}\{([^{}]*)\}', r'(\1/\2)', result)
    for _ in range(3):
        result = re.sub(r'(?<![a-zA-Z])frac\{([^{}]*)\}\{([^{}]*)\}', r'(\1/\2)', result)

    result = re.sub(r'\\sqrt\[([^]]*)\]\{([^}]*)\}', r'\1√(\2)', result)
    result = re.sub(r'\\sqrt\{([^}]*)\}', r'√(\1)', result)

    result = re.sub(r'\\binom\{([^}]*)\}\{([^}]*)\}', r'C(\1,\2)', result)

    for latex, symbol in sorted(SYMBOL_MAP.items(), key=lambda x: -len(x[0])):
        result = result.replace(latex, symbol)
    for latex, func in sorted(TRIG_FUNCS.items(), key=lambda x: -len(x[0])):
        result = result.replace(latex, func)

    for cmd, replacement in BARE_COMMANDS.items():
        result = re.sub(rf'(?<![a-zA-Z\\]){cmd}\{{([^}}]*)\}}', r'\1', result)
        result = re.sub(rf'(?<![a-zA-Z\\]){cmd}(?![a-zA-Z])', replacement, result)

    if use_tags:
        result = re.sub(r'\^\{([^}]*)\}', r'<super>\1</super>', result)
        result = re.sub(r'\^([a-zA-Z0-9°])', r'<super>\1</super>', result)
        result = re.sub(r'_\{([^}]*)\}', r'<sub>\1</sub>', result)
        result = re.sub(r'_([a-zA-Z0-9])', r'<sub>\1</sub>', result)
    else:
        result = re.sub(r'\^\{([^}]*)\}', r'^\1', result)
        result = re.sub(r'\^([a-zA-Z0-9°])', r'^\1', result)
        result = re.sub(r'_\{([^}]*)\}', r'_\1', result)
        result = re.sub(r'_([a-zA-Z0-9])', r'_\1', result)

    result = result.replace('{', '').replace('}', '')
    result = _fix_unicode_scripts(result, use_tags)
    result = re.sub(r'[ \t]+', ' ', result)
    result = re.sub(r' *\n *', '\n', result)
    result = re.sub(r'\n{3,}', '\n\n', result)
    result = result.strip()

    result = _convert_matrix_brackets(result, use_tags)

    return result


def _latex_to_paragraph(text: str) -> str:
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = re.sub(r'\n\s*\n', '__PARABREAK__', text)
    text = text.replace('\n', '__LINEBREAK__')

    result = _process_latex(text, use_tags=True)

    result = result.replace('__PARABREAK__', '<br/><br/>')
    result = result.replace('__LINEBREAK__', '<br/>')

    tags = {}
    for i, tag in enumerate(re.findall(r'</?(?:super|sub|b|i|font[^>]*)>|<br\s*/?>', result)):
        ph = f"__TAG{i}__"
        tags[ph] = tag
        result = result.replace(tag, ph, 1)
    result = result.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    for ph, tag in tags.items():
        result = result.replace(ph, tag)
    return result


def _latex_to_plain(text: str) -> str:
    return _process_latex(text, use_tags=False)


# ═══════════════════════════════════════════════════════════════════════
# Date Formatting
# ═══════════════════════════════════════════════════════════════════════

def _format_date_for_display(paper_date: Optional[str] = None) -> str:
    if paper_date:
        try:
            dt = datetime.strptime(paper_date, "%Y-%m-%d")
            return dt.strftime("%d/%m/%Y")
        except ValueError:
            try:
                dt = datetime.fromisoformat(paper_date.replace('Z', '+00:00'))
                return dt.strftime("%d/%m/%Y")
            except Exception:
                logger.warning(f"Invalid paperDate format: {paper_date}, using today's date")
                return datetime.now().strftime("%d/%m/%Y")
    else:
        return datetime.now().strftime("%d/%m/%Y")


# ═══════════════════════════════════════════════════════════════════════
# Manual Question + Image Helpers
# ═══════════════════════════════════════════════════════════════════════

def _is_manual(q: dict) -> bool:
    return bool(
        q.get("isManual")
        or q.get("is_manual")
        or q.get("validationStatus") == "manual"
        or q.get("validation_status") == "manual"
    )


def _get_image_url(q: dict) -> Optional[str]:
    return q.get("imageUrl") or q.get("image_url") or None


def _get_question_table(q: dict) -> Optional[dict]:
    qt = q.get("questionTable") or q.get("question_table")
    if not qt or not isinstance(qt, dict):
        return None
    if not qt.get("headers") or not qt.get("rows"):
        return None
    return qt


def _fetch_image_bytes(url: str, timeout: int = 8) -> Optional[bytes]:
    if not url or not url.startswith(("http://", "https://")):
        return None
    try:
        req = Request(url, headers={"User-Agent": "A4AI-ExportService/1.0"})
        with urlopen(req, timeout=timeout) as resp:
            if resp.status != 200:
                logger.warning(f"Image fetch returned {resp.status} for {url}")
                return None
            return resp.read()
    except (URLError, HTTPError, TimeoutError) as e:
        logger.warning(f"Failed to fetch image {url}: {e}")
        return None
    except Exception as e:
        logger.warning(f"Unexpected error fetching image {url}: {e}")
        return None


def _render_manual_question_image_pdf(image_url: str, W: float):
    from reportlab.lib.units import cm
    from reportlab.platypus import Image as RLImage, Spacer

    img_bytes = _fetch_image_bytes(image_url)
    if not img_bytes:
        return []

    try:
        img_stream = io.BytesIO(img_bytes)
        max_w = W * 0.80
        max_h = 8 * cm
        img = RLImage(img_stream, width=max_w, height=max_h, kind='proportional')
        img.hAlign = 'CENTER'
        return [Spacer(1, 4), img, Spacer(1, 6)]
    except Exception as e:
        logger.warning(f"Failed to render image in PDF: {e}")
        return []


def _render_manual_question_image_docx(container, image_url: str):
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    img_bytes = _fetch_image_bytes(image_url)
    if not img_bytes:
        return

    try:
        img_stream = io.BytesIO(img_bytes)
        p = container.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_stream, width=Cm(10))
    except Exception as e:
        logger.warning(f"Failed to embed image in DOCX: {e}")


# ═══════════════════════════════════════════════════════════════════════
# Markdown Table Parser
# ═══════════════════════════════════════════════════════════════════════

def _parse_table_row(row_str: str) -> List[str]:
    cells = [c.strip() for c in row_str.split('|')]
    cells = [c for c in cells if c]
    cells = [re.sub(r'\*\*(.+?)\*\*', r'\1', c) for c in cells]
    cells = [re.sub(r'\*(.+?)\*', r'\1', c) for c in cells]
    return cells


def _parse_pipe_table_lines(lines: List[str]) -> tuple:
    headers: List[str] = []
    rows: List[List[str]] = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if re.match(r'^\|?[\s\-:]+\|[\s\-:|]*$', stripped):
            continue
        row = _parse_table_row(stripped)
        if not headers:
            headers = row
        elif row:
            rows.append(row)
    return headers, rows


def _split_text_and_tables(text: str) -> List[dict]:
    if not text:
        return [{'type': 'text', 'content': ''}]

    segments: List[dict] = []

    if '\n' in text:
        lines = text.split('\n')
        i = 0
        current_text: List[str] = []

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            has_pipes = stripped.startswith('|') and stripped.count('|') >= 2
            next_is_sep = (
                i + 1 < len(lines)
                and re.match(r'^\s*\|?[\s\-:]+\|[\s\-:|]*$', lines[i + 1].strip())
            )

            if has_pipes and next_is_sep:
                if current_text:
                    segments.append({'type': 'text', 'content': ' '.join(current_text).strip()})
                    current_text = []
                table_lines = []
                while i < len(lines):
                    tl = lines[i].strip()
                    if tl.startswith('|') or re.match(r'^\|?[\s\-:]+\|', tl):
                        table_lines.append(tl)
                        i += 1
                    else:
                        break
                headers, rows = _parse_pipe_table_lines(table_lines)
                if headers:
                    segments.append({'type': 'table', 'content': (headers, rows)})
            else:
                if stripped:
                    current_text.append(stripped)
                i += 1

        if current_text:
            segments.append({'type': 'text', 'content': ' '.join(current_text).strip()})

    else:
        sep_re = re.compile(r'\s*\|?\s*-{3,}(?:[\|\-\s:]*-{3,})+\s*\|?\s*')
        sep_m = sep_re.search(text)

        if not sep_m:
            return [{'type': 'text', 'content': text}]

        pre_sep = text[:sep_m.start()]
        post_sep = text[sep_m.end():]

        hdr_m = re.search(r'((?:\|[^|]+)+\|)\s*$', pre_sep)
        if not hdr_m:
            return [{'type': 'text', 'content': text}]

        pre_table = pre_sep[:hdr_m.start()].strip()
        headers = _parse_table_row(hdr_m.group(1))

        rows: List[List[str]] = []
        last_end = 0
        for m in re.finditer(r'((?:\|[^|]+)+\|)', post_sep):
            row = _parse_table_row(m.group(1))
            if row:
                rows.append(row)
            last_end = m.end()

        post_table = post_sep[last_end:].strip()

        if pre_table:
            segments.append({'type': 'text', 'content': pre_table})
        if headers and rows:
            segments.append({'type': 'table', 'content': (headers, rows)})
        if post_table:
            segments.append({'type': 'text', 'content': post_table})

    return segments if segments else [{'type': 'text', 'content': text}]


def _strip_markdown_table_from_text(text: str) -> str:
    if not text or '\n' not in text:
        return text

    lines = text.split('\n')
    output_lines = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        has_pipes = stripped.startswith('|') and stripped.count('|') >= 2
        next_is_sep = (
            i + 1 < len(lines)
            and re.match(r'^\s*\|?[\s\-:]+\|[\s\-:|]*$', lines[i + 1].strip())
        )

        if has_pipes and next_is_sep:
            while i < len(lines):
                tl = lines[i].strip()
                if tl.startswith('|') or re.match(r'^\|?[\s\-:]+\|', tl):
                    i += 1
                else:
                    break
        else:
            output_lines.append(line)
            i += 1

    result = '\n'.join(output_lines)
    result = re.sub(r'\n{3,}', '\n\n', result)
    return result.strip()


def _render_inline_table_pdf(headers: List[str], rows: List[List[str]], styles, W: float, tpl: dict) -> list:
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import Table, TableStyle, Paragraph, Spacer

    if not headers or not rows:
        return []

    num_cols = len(headers)
    cell_style = styles.get('Option', styles['Normal'])

    table_data = [[Paragraph(f"<b>{h}</b>", cell_style) for h in headers]]

    for row in rows:
        padded = (row + [''] * num_cols)[:num_cols]
        table_data.append([
            Paragraph(_latex_to_paragraph(str(c)), cell_style)
            for c in padded
        ])

    if num_cols == 2:
        col_widths = [W * 0.60, W * 0.40]
    elif num_cols == 3:
        col_widths = [W * 0.50, W * 0.25, W * 0.25]
    elif num_cols == 4:
        col_widths = [W * 0.40, W * 0.20, W * 0.20, W * 0.20]
    elif num_cols == 5:
        col_widths = [W * 0.10, W * 0.42, W * 0.08, W * 0.20, W * 0.20]
    else:
        col_widths = [W / num_cols] * num_cols

    col_widths = col_widths[:num_cols]

    t = Table(table_data, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle([
        ('GRID',          (0, 0), (-1, -1), 0.5, HexColor(tpl['border'])),
        ('BACKGROUND',    (0, 0), (-1, 0),  HexColor(tpl['table_header_bg'])),
        ('TEXTCOLOR',     (0, 0), (-1, 0),  HexColor(tpl['table_header_text'])),
        ('VALIGN',        (0, 0), (-1, -1), 'TOP'),
        ('TOPPADDING',    (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ('LEFTPADDING',   (0, 0), (-1, -1), 4),
        ('RIGHTPADDING',  (0, 0), (-1, -1), 4),
        ('ALIGN',         (1, 0), (-1, -1), 'RIGHT'),
    ]))
    return [Spacer(1, 4), t, Spacer(1, 6)]


# ═══════════════════════════════════════════════════════════════════════
# Group questions by section
# ═══════════════════════════════════════════════════════════════════════

def _group_by_section(questions: List[dict]) -> dict:
    groups = {}
    for q in questions:
        sec = q.get('section') or q.get('_section') or 'NONE'
        if sec not in groups:
            groups[sec] = []
        groups[sec].append(q)
    return groups


def _has_sections(questions: List[dict]) -> bool:
    for q in questions:
        sec = q.get('section') or q.get('_section')
        if sec and sec in CBSE_SECTIONS_META:
            return True
    return False


def _has_accountancy_sections(questions: List[dict]) -> bool:
    for q in questions:
        sec = q.get('section') or q.get('_section') or ''
        if sec in ACCOUNTANCY_SECTIONS_META or sec.startswith(('A_', 'B1_')):
            return True
    return False


def _get_section_order(questions: List[dict]) -> tuple:
    if _has_accountancy_sections(questions):
        return ACCOUNTANCY_SECTION_ORDER, ACCOUNTANCY_SECTIONS_META
    elif _has_sections(questions):
        return SECTION_ORDER, CBSE_SECTIONS_META
    return None, None


# ═══════════════════════════════════════════════════════════════════════
# Question Table Rendering — PDF  (Statistics — stays semantic blue)
# ═══════════════════════════════════════════════════════════════════════

def _render_question_table_pdf(question_table, styles, W):
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import Table, TableStyle, Paragraph, Spacer

    elements = []
    if not question_table or not isinstance(question_table, dict):
        return elements

    headers = question_table.get("headers", [])
    rows = question_table.get("rows", [])
    caption = question_table.get("caption")

    if not headers or not rows:
        return elements

    num_cols = len(headers)

    if caption:
        elements.append(Spacer(1, 2))
        cap_style = styles.get('QText', styles['Normal'])
        elements.append(Paragraph(
            f"<i>{_latex_to_paragraph(str(caption))}</i>",
            cap_style
        ))
        elements.append(Spacer(1, 3))
    else:
        elements.append(Spacer(1, 4))

    cell_style = styles.get('Option', styles['Normal'])

    table_data = [[Paragraph(f"<b>{_latex_to_paragraph(str(h))}</b>", cell_style) for h in headers]]

    for row in rows:
        if not isinstance(row, list):
            continue
        padded = (row + [""] * num_cols)[:num_cols]
        table_data.append([
            Paragraph(_latex_to_paragraph(str(cell)), cell_style)
            for cell in padded
        ])

    if num_cols == 2:
        col_widths = [W * 0.55, W * 0.45]
    elif num_cols == 3:
        col_widths = [W * 0.40, W * 0.30, W * 0.30]
    elif num_cols == 4:
        col_widths = [W * 0.34, W * 0.22, W * 0.22, W * 0.22]
    elif num_cols == 5:
        col_widths = [W * 0.28, W * 0.18, W * 0.18, W * 0.18, W * 0.18]
    else:
        col_widths = [W / num_cols] * num_cols

    col_widths = col_widths[:num_cols]

    t = Table(table_data, colWidths=col_widths, repeatRows=1, hAlign='CENTER')

    style_cmds = [
        ('GRID',          (0, 0), (-1, -1), 0.5, HexColor('#bfdbfe')),
        ('BACKGROUND',    (0, 0), (-1, 0),  HexColor('#dbeafe')),
        ('TEXTCOLOR',     (0, 0), (-1, 0),  HexColor('#1e3a8a')),
        ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
        ('ALIGN',         (0, 0), (-1, 0),  'CENTER'),
        ('TOPPADDING',    (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ('LEFTPADDING',   (0, 0), (-1, -1), 6),
        ('RIGHTPADDING',  (0, 0), (-1, -1), 6),
        ('ALIGN',         (1, 1), (-1, -1), 'RIGHT'),
        ('ALIGN',         (0, 1), (0, -1),  'LEFT'),
    ]

    t.setStyle(TableStyle(style_cmds))
    elements.append(t)
    elements.append(Spacer(1, 8))

    return elements


# ═══════════════════════════════════════════════════════════════════════
# Question Table Rendering — DOCX (Statistics — stays semantic blue)
# ═══════════════════════════════════════════════════════════════════════

def _render_question_table_docx(container, question_table):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    if not question_table or not isinstance(question_table, dict):
        return

    headers = question_table.get("headers", [])
    rows = question_table.get("rows", [])
    caption = question_table.get("caption")

    if not headers or not rows:
        return

    num_cols = len(headers)

    if caption:
        cp = container.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(_latex_to_plain(str(caption)))
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = RGBColor(75, 85, 99)

    total_data_rows = 1 + len(rows)

    table = container.add_table(rows=total_data_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(_latex_to_plain(str(header)))
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(30, 58, 138)

        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): 'DBEAFE',
            qn('w:val'): 'clear',
        })
        shading.append(shd)

    for i, row in enumerate(rows):
        if not isinstance(row, list):
            continue
        padded = (row + [""] * num_cols)[:num_cols]
        for j, cell_text in enumerate(padded):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(_latex_to_plain(str(cell_text)))
            r.font.size = Pt(10)

            if j > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    container.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════
# Accountancy Answer Table Rendering — PDF (stays semantic emerald)
# ═══════════════════════════════════════════════════════════════════════

def _render_answer_table_pdf(answer_table, styles, W):
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import Table, TableStyle, Paragraph, Spacer

    elements = []
    if not answer_table or not isinstance(answer_table, dict):
        return elements

    table_type = answer_table.get("type", "")
    headers = answer_table.get("headers", [])
    rows = answer_table.get("rows", [])
    total_row = answer_table.get("total_row")

    if not headers or not rows:
        return elements

    num_cols = len(headers)

    title_map = {
        "journal_entry": "Journal Entry",
        "ledger": "Ledger Account",
        "trial_balance": "Trial Balance",
    }
    title = title_map.get(table_type, "Answer Table")
    elements.append(Spacer(1, 6))
    elements.append(Paragraph(
        f"<b>{title}:</b>",
        styles.get('AnswerLine', styles['Normal'])
    ))
    elements.append(Spacer(1, 4))

    header_style = styles.get('Option', styles['Normal'])
    table_data = [[Paragraph(f"<b>{h}</b>", header_style) for h in headers]]

    for row in rows:
        if not isinstance(row, list):
            continue
        padded = (row + [""] * num_cols)[:num_cols]
        table_data.append([
            Paragraph(_latex_to_paragraph(str(cell)), styles.get('Option', styles['Normal']))
            for cell in padded
        ])

    if total_row and isinstance(total_row, list):
        padded_total = (total_row + [""] * num_cols)[:num_cols]
        table_data.append([
            Paragraph(f"<b>{_latex_to_paragraph(str(cell))}</b>", styles.get('Option', styles['Normal']))
            for cell in padded_total
        ])

    if table_type == "journal_entry":
        col_widths = [W * 0.13, W * 0.40, W * 0.07, W * 0.20, W * 0.20]
    elif table_type == "ledger":
        col_w = W / 8
        col_widths = [col_w] * num_cols
    elif table_type == "trial_balance":
        col_widths = [W * 0.08, W * 0.42, W * 0.10, W * 0.20, W * 0.20]
    else:
        col_widths = [W / num_cols] * num_cols

    col_widths = col_widths[:num_cols]

    t = Table(table_data, colWidths=col_widths, repeatRows=1)

    style_cmds = [
        ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#d1d5db')),
        ('BACKGROUND', (0, 0), (-1, 0), HexColor('#f3f4f6')),
        ('TEXTCOLOR', (0, 0), (-1, 0), HexColor('#1f2937')),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
    ]

    if table_type == "journal_entry":
        style_cmds.append(('ALIGN', (3, 0), (4, -1), 'RIGHT'))
    elif table_type == "trial_balance":
        style_cmds.append(('ALIGN', (3, 0), (4, -1), 'RIGHT'))
    elif table_type == "ledger":
        if num_cols >= 8:
            style_cmds.append(('ALIGN', (3, 0), (3, -1), 'RIGHT'))
            style_cmds.append(('ALIGN', (7, 0), (7, -1), 'RIGHT'))
            style_cmds.append(('LINEAFTER', (3, 0), (3, -1), 1.5, HexColor('#374151')))

    if total_row:
        last_row = len(table_data) - 1
        style_cmds.extend([
            ('BACKGROUND', (0, last_row), (-1, last_row), HexColor('#e5e7eb')),
            ('LINEABOVE', (0, last_row), (-1, last_row), 1.5, HexColor('#374151')),
        ])

    t.setStyle(TableStyle(style_cmds))
    elements.append(t)
    elements.append(Spacer(1, 6))

    return elements


# ═══════════════════════════════════════════════════════════════════════
# Accountancy Answer Table Rendering — DOCX (stays semantic emerald)
# ═══════════════════════════════════════════════════════════════════════

def _render_answer_table_docx(container, answer_table):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    if not answer_table or not isinstance(answer_table, dict):
        return

    table_type = answer_table.get("type", "")
    headers = answer_table.get("headers", [])
    rows = answer_table.get("rows", [])
    total_row = answer_table.get("total_row")

    if not headers or not rows:
        return

    num_cols = len(headers)

    title_map = {
        "journal_entry": "Journal Entry",
        "ledger": "Ledger Account",
        "trial_balance": "Trial Balance",
    }
    title = title_map.get(table_type, "Answer Table")
    tp = container.add_paragraph()
    tr = tp.add_run(f"{title}:")
    tr.bold = True
    tr.font.size = Pt(10)
    tr.font.color.rgb = RGBColor(4, 120, 87)

    total_data_rows = 1 + len(rows) + (1 if total_row else 0)

    table = container.add_table(rows=total_data_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(header)
        r.bold = True
        r.font.size = Pt(9)
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): 'F3F4F6',
            qn('w:val'): 'clear',
        })
        shading.append(shd)

    for i, row in enumerate(rows):
        if not isinstance(row, list):
            continue
        padded = (row + [""] * num_cols)[:num_cols]
        for j, cell_text in enumerate(padded):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(cell_text))
            r.font.size = Pt(9)

            if table_type in ("journal_entry", "trial_balance") and j >= num_cols - 2:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif table_type == "ledger" and num_cols >= 8 and j in (3, 7):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if total_row and isinstance(total_row, list):
        row_idx = 1 + len(rows)
        padded_total = (total_row + [""] * num_cols)[:num_cols]
        for j, cell_text in enumerate(padded_total):
            cell = table.rows[row_idx].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(cell_text))
            r.bold = True
            r.font.size = Pt(9)

            if table_type in ("journal_entry", "trial_balance") and j >= num_cols - 2:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            shading = cell._element.get_or_add_tcPr()
            shd = shading.makeelement(qn('w:shd'), {
                qn('w:fill'): 'E5E7EB',
                qn('w:val'): 'clear',
            })
            shading.append(shd)

    container.add_paragraph()


def _render_inline_table_docx(container, headers: List[str], rows: List[List[str]]):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    if not headers or not rows:
        return

    num_cols = len(headers)
    total_rows = 1 + len(rows)

    table = container.add_table(rows=total_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)

    for i, row in enumerate(rows):
        padded = (row + [''] * num_cols)[:num_cols]
        for j, val in enumerate(padded):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            if j > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    container.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════
# Question Card Wrapper (PDF) — card_based layout, template-aware
# ═══════════════════════════════════════════════════════════════════════

def _wrap_question_card(elements, W, tpl: dict, section_label: Optional[str] = None):
    from reportlab.platypus import Table, TableStyle, Spacer
    from reportlab.lib.colors import HexColor

    clean = list(elements)
    while clean and isinstance(clean[-1], Spacer):
        clean.pop()

    if not clean:
        return [Spacer(1, _sc(4, tpl))]

    style = tpl.get("card_style", "card")

    if style == "flat":
        return clean + [Spacer(1, _sc(8, tpl))]

    t = Table([[clean]], colWidths=[W])

    style_cmds = [
        ('BACKGROUND',    (0, 0), (-1, -1), HexColor(tpl['card_bg'])),
        ('BOX',           (0, 0), (-1, -1), 0.5, HexColor(tpl['card_border'])),
        ('TOPPADDING',    (0, 0), (-1, -1), _sc(5, tpl)),
        ('BOTTOMPADDING', (0, 0), (-1, -1), _sc(5, tpl)),
        ('LEFTPADDING',   (0, 0), (-1, -1), _sc(6, tpl)),
        ('RIGHTPADDING',  (0, 0), (-1, -1), _sc(6, tpl)),
    ]

    if style == "stripe":
        accent = _section_color(tpl, section_label)
        style_cmds.append(('LINEBEFORE', (0, 0), (0, -1), 3, HexColor(accent)))
        style_cmds.append(('LEFTPADDING', (0, 0), (-1, -1), _sc(10, tpl)))

    try:
        style_cmds.append(('ROUNDEDCORNERS', [6, 6, 6, 6]))
    except Exception:
        pass

    t.setStyle(TableStyle(style_cmds))

    return [t, Spacer(1, _sc(4, tpl))]


def _render_or_separator(styles, W, tpl: dict):
    from reportlab.platypus import Paragraph, Spacer, HRFlowable, Table
    from reportlab.lib.colors import HexColor

    or_elements = [
        Spacer(1, 2),
        Table(
            [[
                HRFlowable(width="30%", thickness=0.5, color=HexColor(tpl['border'])),
                Paragraph("<b>OR</b>", styles.get('SectionHeader', styles['Normal'])),
                HRFlowable(width="30%", thickness=0.5, color=HexColor(tpl['border'])),
            ]],
            colWidths=[W * 0.35, W * 0.30, W * 0.35],
        ),
        Spacer(1, 2),
    ]
    return or_elements


# ═══════════════════════════════════════════════════════════════════════
# PDF Generation  (entry point — routes to card_based or institute_paper)
# ═══════════════════════════════════════════════════════════════════════

def generate_pdf(
    questions: List[dict],
    exam_title: str = "Test Paper",
    board: str = "CBSE",
    class_grade: str = "10",
    subject: str = "Science",
    include_answers: bool = False,
    include_explanations: bool = False,
    logo_base64: Optional[str] = None,
    paper_date: Optional[str] = None,
    template: str = DEFAULT_TEMPLATE,
    teacher_name: Optional[str] = None,
    institute_name: Optional[str] = None,
    duration: Optional[str] = None,
    topic: Optional[str] = None,
) -> bytes:
    tpl = _get_template(template)

    # v14: colorful -> institute-paper layout
    if tpl.get("layout_style") == "institute_paper":
        return _generate_pdf_institute(
            questions=questions, exam_title=exam_title, board=board,
            class_grade=class_grade, subject=subject,
            include_answers=include_answers, include_explanations=include_explanations,
            logo_base64=logo_base64, paper_date=paper_date, tpl=tpl,
            teacher_name=teacher_name, institute_name=institute_name,
            duration=duration, topic=topic,
        )

    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, HRFlowable, Image as RLImage,
    )

    buffer = io.BytesIO()

    top_m, bottom_m, left_m, right_m = tpl['margins_cm']
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        topMargin=top_m * cm, bottomMargin=bottom_m * cm,
        leftMargin=left_m * cm, rightMargin=right_m * cm,
    )

    styles = getSampleStyleSheet()

    W = A4[0] - (left_m + right_m) * cm

    fb = tpl['font_body']
    fbd = tpl['font_bold']

    custom_styles = {
        'SchoolName': dict(parent=styles['Title'], fontSize=14, leading=18, spaceAfter=2, alignment=TA_CENTER, textColor=HexColor(tpl['primary']), fontName=fbd),
        'ExamMeta': dict(parent=styles['Normal'], fontSize=10, alignment=TA_CENTER, textColor=HexColor(tpl['secondary']), spaceAfter=4, fontName=fb),
        'SectionHeader': dict(parent=styles['Heading1'], fontSize=11, spaceBefore=_sc(10, tpl), spaceAfter=2, textColor=HexColor(tpl['primary']), fontName=fbd, alignment=TA_CENTER),
        'SectionSub': dict(parent=styles['Normal'], fontSize=9, alignment=TA_CENTER, textColor=HexColor(tpl['muted']), spaceAfter=2, fontName=fb),
        'SectionInstruction': dict(parent=styles['Normal'], fontSize=8.5, alignment=TA_CENTER, textColor=HexColor(tpl['light_muted']), spaceAfter=4, leading=10, fontName=fb),
        'SectionTitle': dict(parent=styles['Heading2'], fontSize=11, spaceBefore=_sc(14, tpl), spaceAfter=6, textColor=HexColor(tpl['primary']), fontName=fbd),
        'QText': dict(parent=styles['Normal'], fontSize=10, spaceBefore=2, spaceAfter=1, leading=12, textColor=HexColor('#1f1f3a'), fontName=fb),
        'Option': dict(parent=styles['Normal'], fontSize=9.5, leftIndent=14, spaceBefore=1, spaceAfter=1, leading=11.5, textColor=HexColor('#333355'), fontName=fb),
        'CorrectOption': dict(parent=styles['Normal'], fontSize=9.5, leftIndent=14, spaceBefore=1, spaceAfter=1, leading=11.5, textColor=HexColor(tpl['correct']), fontName=fbd),
        'AnswerLine': dict(parent=styles['Normal'], fontSize=9, leftIndent=14, spaceBefore=1, textColor=HexColor(tpl['correct']), fontName=fbd),
        'Explanation': dict(parent=styles['Normal'], fontSize=8.5, leftIndent=14, spaceBefore=1, spaceAfter=3, textColor=HexColor(tpl['muted']), leading=11, fontName=fb),
        'Marks': dict(parent=styles['Normal'], fontSize=9, alignment=TA_RIGHT, textColor=HexColor(tpl['light_muted']), fontName=fb),
        'Instruction': dict(parent=styles['Normal'], fontSize=9, leftIndent=12, spaceBefore=2, spaceAfter=2, textColor=HexColor(tpl['secondary']), leading=12, fontName=fb),
        'FooterText': dict(parent=styles['Normal'], fontSize=8, textColor=HexColor(tpl['light_muted']), alignment=TA_CENTER, fontName=fb),
        'ORText': dict(parent=styles['Normal'], fontSize=10, alignment=TA_CENTER, textColor=HexColor(tpl['muted']), fontName=fbd, spaceBefore=2, spaceAfter=2),
        'ManualBadge': dict(parent=styles['Normal'], fontSize=7.5, textColor=HexColor('#4f46e5'), fontName=fbd),
    }
    for name, props in custom_styles.items():
        try:
            styles.add(ParagraphStyle(name=name, **props))
        except KeyError:
            pass

    story = []

    display_date = _format_date_for_display(paper_date)

    logo_img = None
    if logo_base64:
        try:
            if ',' in logo_base64:
                logo_base64 = logo_base64.split(',', 1)[1]
            logo_img = RLImage(io.BytesIO(base64.b64decode(logo_base64)), width=1.8 * cm, height=1.8 * cm)
            logo_img.hAlign = 'CENTER'
        except Exception as e:
            logger.warning(f"Logo failed: {e}")

    total_marks = sum(q.get('marks', 1) for q in questions)

    title_block = [
        Paragraph(f"<b>{exam_title}</b>", styles['SchoolName']),
        Paragraph(f"{board} Board | Class {class_grade} | {subject}", styles['ExamMeta']),
    ]
    info_block = [
        Paragraph(f"Date: {display_date}", styles['ExamMeta']),
        Paragraph(f"Total Marks: {total_marks}", styles['ExamMeta']),
        Paragraph(f"Total Questions: {len(questions)}", styles['ExamMeta']),
    ]

    if logo_img:
        ht = Table([[logo_img, title_block, info_block]], colWidths=[2.5 * cm, W - 6 * cm, 3.5 * cm])
    else:
        ht = Table([[title_block, info_block]], colWidths=[W - 4 * cm, 4 * cm])
    ht.setStyle(TableStyle([('VALIGN', (0, 0), (-1, -1), 'MIDDLE'), ('TOPPADDING', (0, 0), (-1, -1), 4), ('BOTTOMPADDING', (0, 0), (-1, -1), 4)]))

    if tpl['header_style'] == 'banner':
        banner = Table([[ht]], colWidths=[W])
        banner.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), HexColor(tpl.get('banner_bg', '#F5F3FF'))),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        ]))
        story.append(banner)
        rule_thickness = 2.5
    elif tpl['header_style'] == 'minimal':
        story.append(ht)
        rule_thickness = 0.75
    else:
        story.append(ht)
        rule_thickness = 1.5

    story.append(Spacer(1, _sc(4, tpl)))
    story.append(HRFlowable(width="100%", thickness=rule_thickness, color=HexColor(tpl['primary']), spaceAfter=_sc(8, tpl)))

    sec_order, sec_meta_dict = _get_section_order(questions)
    has_sec = sec_order is not None

    story.append(Paragraph("<b>General Instructions:</b>", styles['SectionTitle']))

    base_instructions = [
        "All questions are compulsory.",
        "Read each question carefully before answering.",
    ]

    if sec_meta_dict is ACCOUNTANCY_SECTIONS_META:
        base_instructions.extend([
            "This question paper is divided into <b>Part A</b> and <b>Part B</b>.",
            "<b>Part A</b> is compulsory for all candidates.",
            "<b>Part B</b> has two options — attempt only one.",
            "Internal choice has been provided in some questions.",
        ])
    elif has_sec:
        base_instructions.extend([
            f"This question paper has <b>5 Sections</b> — A, B, C, D, and E.",
            f"<b>Section A</b> has 20 MCQs / Assertion-Reason (1 mark each).",
            f"<b>Section B</b> has 5 Very Short Answer questions (2 marks each).",
            f"<b>Section C</b> has 6 Short Answer questions (3 marks each).",
            f"<b>Section D</b> has 4 Long Answer questions (5 marks each).",
            f"<b>Section E</b> has 3 Case Study questions (4 marks each).",
        ])
    else:
        base_instructions.extend([
            "For MCQs, select the <b>best answer</b> from the given choices.",
        ])

    base_instructions.append(f"Total marks: <b>{total_marks}</b>. Time allotted as per school schedule.")

    for inst in base_instructions:
        story.append(Paragraph(f"• {inst}", styles['Instruction']))
    story.append(Spacer(1, _sc(6, tpl)))
    story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor(tpl['border']), spaceAfter=_sc(6, tpl)))

    labels = ["A", "B", "C", "D", "E", "F"]

    QW = W - 20

    def _render_question(q, q_num):
        from reportlab.platypus import Table as RLTable, TableStyle as RLTableStyle
        elements = []
        raw_text = q.get('text', '')
        marks = q.get('marks', 1)
        marks_label = f"[{marks} {'mark' if marks == 1 else 'marks'}]"

        question_table = _get_question_table(q)

        if question_table:
            raw_text = _strip_markdown_table_from_text(raw_text)

        segments = _split_text_and_tables(raw_text)

        first_text = ''
        for seg in segments:
            if seg['type'] == 'text' and seg['content']:
                first_text = _latex_to_paragraph(seg['content'])
                break
        if not first_text:
            first_text = _latex_to_paragraph(raw_text)

        q_text_html = f"<b>Q{q_num}.</b> {first_text}"

        qt = RLTable(
            [[Paragraph(q_text_html, styles['QText']),
              Paragraph(marks_label, styles['Marks'])]],
            colWidths=[QW * 0.84, QW * 0.16],
        )
        qt.setStyle(RLTableStyle([
            ('VALIGN',        (0, 0), (-1, -1), 'TOP'),
            ('TOPPADDING',    (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ]))
        elements.append(qt)

        if question_table:
            qt_elements = _render_question_table_pdf(question_table, styles, QW)
            elements.extend(qt_elements)

        image_url = _get_image_url(q)
        if image_url:
            img_elements = _render_manual_question_image_pdf(image_url, QW)
            elements.extend(img_elements)

        first_text_skipped = False
        for seg in segments:
            if seg['type'] == 'text':
                if not first_text_skipped:
                    first_text_skipped = True
                    continue
                content = _latex_to_paragraph(seg['content'])
                if content:
                    elements.append(Paragraph(content, styles['QText']))
            elif seg['type'] == 'table':
                if question_table:
                    continue
                hdrs, rws = seg['content']
                elements.extend(_render_inline_table_pdf(hdrs, rws, styles, QW, tpl))

        options = q.get('options', [])
        correct_answer = q.get('correctAnswer', q.get('correct_answer', ''))

        if options:
            for opt_idx, opt in enumerate(options):
                opt_text = _latex_to_paragraph(opt)
                letter = labels[opt_idx] if opt_idx < len(labels) else str(opt_idx + 1)
                is_correct = False
                if include_answers and correct_answer:
                    ca = correct_answer.strip()
                    if ca.upper().startswith(letter) or opt.strip() == ca.strip():
                        is_correct = True
                style = styles['CorrectOption'] if is_correct else styles['Option']
                opt_clean = re.sub(r'^[A-F][).\s]+\s*', '', opt_text).strip()
                prefix = f"<b>{letter})</b> " if is_correct else f"{letter}) "
                elements.append(Paragraph(f"{prefix}{opt_clean}", style))
        else:
            fmt = q.get('format', 'mcq')
            if not include_answers:
                if fmt == 'short_answer':
                    elements.append(Spacer(1, _sc(18, tpl)))
                elif fmt == 'long_answer':
                    elements.append(Spacer(1, _sc(40, tpl)))
                elif fmt in ('journal_entry', 'ledger', 'trial_balance'):
                    elements.append(Spacer(1, _sc(50, tpl)))
                elif fmt == 'image':
                    elements.append(Spacer(1, _sc(20, tpl)))

        if include_answers and include_explanations:
            raw_table = q.get('answer_table') or q.get('answerTable')
            if raw_table and isinstance(raw_table, dict):
                elements.extend(_render_answer_table_pdf(raw_table, styles, QW))
            else:
                ans = _latex_to_paragraph(correct_answer)
                elements.append(Paragraph(f"<b>Answer:</b> {ans}", styles['AnswerLine']))

        if include_explanations:
            exp = _latex_to_paragraph(q.get('explanation', ''))
            if exp:
                elements.append(Paragraph(f"<b>Explanation:</b> {exp}", styles['Explanation']))

        return elements

    q_num = 0

    if has_sec:
        grouped = _group_by_section(questions)
        last_section_title = None

        for sec_key in sec_order:
            sec_qs = grouped.get(sec_key, [])
            if not sec_qs:
                continue

            meta = sec_meta_dict.get(sec_key, {})
            current_title = meta.get('title', sec_key)
            sec_letter = sec_key[:1]
            sec_color = _section_color(tpl, sec_letter)

            if current_title != last_section_title:
                story.append(HRFlowable(width="60%", thickness=1, color=HexColor(sec_color), spaceBefore=_sc(14, tpl), spaceAfter=_sc(4, tpl)))
                story.append(Paragraph(f'<font color="{sec_color}"><b>{current_title}</b></font>', styles['SectionHeader']))
                last_section_title = current_title

            story.append(Paragraph(meta.get('subtitle', ''), styles['SectionSub']))
            story.append(Paragraph(meta.get('instruction', ''), styles['SectionInstruction']))
            story.append(HRFlowable(width="40%", thickness=0.5, color=HexColor(tpl['border']), spaceAfter=_sc(6, tpl)))

            main_qs = [q for q in sec_qs if not q.get('_is_or', False)]
            or_qs = [q for q in sec_qs if q.get('_is_or', False)]
            or_queue = list(or_qs)

            for q in main_qs:
                q_num += 1
                elements = _render_question(q, q_num)
                card = _wrap_question_card(elements, W, tpl, sec_letter)
                story.extend(card)

                if or_queue:
                    or_q = or_queue.pop(0)
                    story.extend(_render_or_separator(styles, W, tpl))
                    or_elements = _render_question(or_q, q_num)
                    or_card = _wrap_question_card(or_elements, W, tpl, sec_letter)
                    story.extend(or_card)

            for or_q in or_queue:
                q_num += 1
                story.extend(_render_or_separator(styles, W, tpl))
                or_elements = _render_question(or_q, q_num)
                or_card = _wrap_question_card(or_elements, W, tpl, sec_letter)
                story.extend(or_card)

        unsectioned = grouped.get('NONE', [])
        if unsectioned:
            story.append(HRFlowable(width="60%", thickness=1, color=HexColor('#4f46e5'), spaceBefore=_sc(14, tpl), spaceAfter=_sc(4, tpl)))
            story.append(Paragraph(f"<b>Additional Questions</b>", styles['SectionHeader']))
            story.append(Paragraph("(Added by teacher)", styles['SectionSub']))
            story.append(HRFlowable(width="40%", thickness=0.5, color=HexColor(tpl['border']), spaceAfter=_sc(6, tpl)))

            for q in unsectioned:
                q_num += 1
                elements = _render_question(q, q_num)
                card = _wrap_question_card(elements, W, tpl, None)
                story.extend(card)

    else:
        for q in questions:
            q_num += 1
            elements = _render_question(q, q_num)
            card = _wrap_question_card(elements, W, tpl, None)
            story.extend(card)

    if include_answers and not include_explanations:
        story.append(PageBreak())
        story.append(Paragraph("<b>Answer Key</b>", styles['SchoolName']))
        story.append(HRFlowable(width="100%", thickness=1, color=HexColor(tpl['primary']), spaceAfter=_sc(10, tpl)))

        q_num_ak = 0
        all_qs_ordered = []
        if has_sec:
            grouped = _group_by_section(questions)
            for sec_key in sec_order:
                all_qs_ordered.extend(grouped.get(sec_key, []))
            all_qs_ordered.extend(grouped.get('NONE', []))
        else:
            all_qs_ordered = questions

        for q in all_qs_ordered:
            q_num_ak += 1
            raw_table = q.get('answer_table') or q.get('answerTable')
            if raw_table and isinstance(raw_table, dict):
                story.append(Paragraph(f"<b>Q{q_num_ak}.</b>", styles['QText']))
                table_elements = _render_answer_table_pdf(raw_table, styles, W)
                story.extend(table_elements)
                story.append(Spacer(1, 4))
            else:
                correct = _latex_to_paragraph(q.get('correctAnswer', q.get('correct_answer', '')))
                story.append(Paragraph(f"<b>Q{q_num_ak}.</b> {correct}", styles['QText']))
                story.append(Spacer(1, 2))

    story.append(Spacer(1, 20))
    story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor(tpl['border']), spaceAfter=6))
    story.append(Paragraph(f"Generated by a4ai · {board} {subject} Class {class_grade} · {display_date}", styles['FooterText']))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


# ═══════════════════════════════════════════════════════════════════════
# DOCX Generation  (entry point — routes to card_based or institute_paper)
# ═══════════════════════════════════════════════════════════════════════

def generate_docx(
    questions: List[dict],
    exam_title: str = "Test Paper",
    board: str = "CBSE",
    class_grade: str = "10",
    subject: str = "Science",
    include_answers: bool = False,
    include_explanations: bool = False,
    logo_base64: Optional[str] = None,
    paper_date: Optional[str] = None,
    template: str = DEFAULT_TEMPLATE,
    teacher_name: Optional[str] = None,
    institute_name: Optional[str] = None,
    duration: Optional[str] = None,
    topic: Optional[str] = None,
) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    tpl = _get_template(template)

    # v14: colorful -> institute-paper layout
    if tpl.get("layout_style") == "institute_paper":
        return _generate_docx_institute(
            questions=questions, exam_title=exam_title, board=board,
            class_grade=class_grade, subject=subject,
            include_answers=include_answers, include_explanations=include_explanations,
            logo_base64=logo_base64, paper_date=paper_date, tpl=tpl,
            teacher_name=teacher_name, institute_name=institute_name,
            duration=duration, topic=topic,
        )

    doc = Document()

    try:
        normal_style = doc.styles['Normal']
        normal_style.font.name = tpl['docx_font']
        rpr = normal_style.element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = rpr.makeelement(qn('w:rFonts'), {})
            rpr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), tpl['docx_font'])
    except Exception as e:
        logger.warning(f"Could not set default docx font: {e}")

    top_m, bottom_m, left_m, right_m = tpl['margins_cm']
    for section in doc.sections:
        section.top_margin = Cm(top_m)
        section.bottom_margin = Cm(bottom_m)
        section.left_margin = Cm(left_m)
        section.right_margin = Cm(right_m)

    def _spt(val: float) -> "Pt":
        return Pt(max(1, val * tpl.get('spacing_scale', 1.0)))

    display_date = _format_date_for_display(paper_date)

    header_container = doc
    if tpl['header_style'] == 'banner':
        banner_table = doc.add_table(rows=1, cols=1)
        banner_cell = banner_table.rows[0].cells[0]
        tcPr = banner_cell._element.get_or_add_tcPr()
        shd = tcPr.makeelement(qn('w:shd'), {qn('w:fill'): _hexnc(tpl.get('banner_bg', '#EEF2FF')), qn('w:val'): 'clear'})
        tcPr.append(shd)
        header_container = banner_cell

    if logo_base64:
        try:
            if ',' in logo_base64:
                logo_base64 = logo_base64.split(',', 1)[1]
            p = header_container.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(io.BytesIO(base64.b64decode(logo_base64)), width=Cm(2))
        except Exception as e:
            logger.warning(f"DOCX logo failed: {e}")

    title_p = header_container.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(exam_title or "Test Paper")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = _rgb(tpl['primary'])
    title_run.font.name = tpl['docx_font']

    sub = header_container.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"{board} Board | Class {class_grade} | {subject}")
    r.font.size = Pt(11)
    r.font.color.rgb = _rgb(tpl['secondary'])
    r.font.name = tpl['docx_font']

    total_marks = sum(q.get('marks', 1) for q in questions)

    meta = header_container.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(f"Total Questions: {len(questions)} | Total Marks: {total_marks} | Date: {display_date}")
    r.font.size = Pt(9)
    r.font.color.rgb = _rgb(tpl['muted'])
    r.font.name = tpl['docx_font']

    rule_char = "─" if tpl['header_style'] == 'minimal' else "━"
    rule_len = 36 if tpl['spacing_scale'] < 1.0 else 50
    doc.add_paragraph(rule_char * rule_len)

    sec_order, sec_meta_dict = _get_section_order(questions)
    has_sec = sec_order is not None

    doc.add_heading("General Instructions", level=2)

    instructions = ["All questions are compulsory.", "Read each question carefully."]
    if sec_meta_dict is ACCOUNTANCY_SECTIONS_META:
        instructions.extend([
            "This paper is divided into Part A and Part B.",
            "Part A is compulsory for all candidates.",
            "Part B has two options — attempt only one.",
            "Internal choice has been provided in some questions.",
        ])
    elif has_sec:
        instructions.extend([
            "This paper has 5 Sections — A, B, C, D, and E.",
            "Section A: 20 questions × 1 mark (MCQ / Assertion-Reason)",
            "Section B: 5 questions × 2 marks (Very Short Answer)",
            "Section C: 6 questions × 3 marks (Short Answer)",
            "Section D: 4 questions × 5 marks (Long Answer)",
            "Section E: 3 questions × 4 marks (Case Study Based)",
        ])
    else:
        instructions.append("For MCQs, select the best answer.")
    instructions.append(f"Total marks: {total_marks}.")

    for inst in instructions:
        p = doc.add_paragraph(inst, style='List Bullet')
        p.paragraph_format.space_after = _spt(2)
    doc.add_paragraph(rule_char * rule_len)

    labels = ["A", "B", "C", "D", "E", "F"]
    q_num = 0

    def _add_docx_separator(container):
        sep_p = container.add_paragraph()
        sep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sep_p.paragraph_format.space_before = _spt(2)
        sep_p.paragraph_format.space_after = _spt(2)
        r = sep_p.add_run("─" * 60)
        r.font.size = Pt(6)
        r.font.color.rgb = _rgb(tpl['border'])

    def _add_docx_or_separator(container):
        sep_p = container.add_paragraph()
        sep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sep_p.paragraph_format.space_before = _spt(4)
        sep_p.paragraph_format.space_after = _spt(4)
        r = sep_p.add_run("─── OR ───")
        r.font.size = Pt(10)
        r.font.color.rgb = _rgb(tpl['muted'])
        r.bold = True

    def _start_question_container(section_letter=None):
        style = tpl.get('card_style', 'card')
        if style == 'flat':
            return doc, None

        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        tcPr = cell._element.get_or_add_tcPr()

        shd = tcPr.makeelement(qn('w:shd'), {qn('w:fill'): _hexnc(tpl['card_bg']), qn('w:val'): 'clear'})
        tcPr.append(shd)

        tcBorders = tcPr.makeelement(qn('w:tcBorders'), {})
        if style == 'stripe':
            accent = _hexnc(_section_color(tpl, section_letter))
            left = tcBorders.makeelement(qn('w:left'), {qn('w:val'): 'single', qn('w:sz'): '24', qn('w:color'): accent})
            tcBorders.append(left)
        else:
            border_hex = _hexnc(tpl['card_border'])
            for side in ('top', 'left', 'bottom', 'right'):
                b = tcBorders.makeelement(qn(f'w:{side}'), {qn('w:val'): 'single', qn('w:sz'): '6', qn('w:color'): border_hex})
                tcBorders.append(b)
        tcPr.append(tcBorders)

        return cell, table

    def _render_q_docx(q, q_num, section_letter=None):
        container, wrapper_table = _start_question_container(section_letter)

        raw_text = q.get('text', '')
        marks = q.get('marks', 1)

        question_table = _get_question_table(q)

        if question_table:
            raw_text = _strip_markdown_table_from_text(raw_text)

        segments = _split_text_and_tables(raw_text)
        first_text = ''
        for seg in segments:
            if seg['type'] == 'text' and seg['content']:
                first_text = _latex_to_plain(seg['content'])
                break
        if not first_text:
            first_text = _latex_to_plain(raw_text)

        p = container.add_paragraph()
        rq = p.add_run(f"Q{q_num}. ")
        rq.bold = True
        rq.font.size = Pt(11)
        rq.font.name = tpl['docx_font']
        rt = p.add_run(first_text)
        rt.font.size = Pt(11)
        rt.font.name = tpl['docx_font']

        rm = p.add_run(f"  [{marks} {'mark' if marks == 1 else 'marks'}]")
        rm.font.size = Pt(8)
        rm.font.color.rgb = _rgb(tpl['light_muted'])
        rm.font.name = tpl['docx_font']

        if question_table:
            _render_question_table_docx(container, question_table)

        image_url = _get_image_url(q)
        if image_url:
            _render_manual_question_image_docx(container, image_url)

        first_text_skipped = False
        for seg in segments:
            if seg['type'] == 'text':
                if not first_text_skipped:
                    first_text_skipped = True
                    continue
                content = _latex_to_plain(seg['content'])
                if content:
                    cp = container.add_paragraph()
                    crun = cp.add_run(content)
                    crun.font.size = Pt(11)
                    crun.font.name = tpl['docx_font']
            elif seg['type'] == 'table':
                if question_table:
                    continue
                hdrs, rws = seg['content']
                _render_inline_table_docx(container, hdrs, rws)

        options = q.get('options', [])
        correct_answer = q.get('correctAnswer', q.get('correct_answer', ''))

        for opt_idx, opt in enumerate(options):
            opt_clean = _latex_to_plain(opt)
            letter = labels[opt_idx] if opt_idx < len(labels) else str(opt_idx + 1)
            is_correct = include_answers and correct_answer and correct_answer.strip().upper().startswith(letter)

            op = container.add_paragraph()
            op.paragraph_format.left_indent = Pt(24)
            op.paragraph_format.space_after = _spt(2)
            opt_stripped = re.sub(r'^[A-F][).\s]+\s*', '', opt_clean).strip()
            run = op.add_run(f"{letter}) {opt_stripped}")
            run.font.size = Pt(10)
            run.font.name = tpl['docx_font']
            if is_correct:
                run.bold = True
                run.font.color.rgb = _rgb(tpl['correct'])

        if include_answers and include_explanations:
            raw_table = q.get('answer_table') or q.get('answerTable')
            if raw_table and isinstance(raw_table, dict):
                _render_answer_table_docx(container, raw_table)
            else:
                correct = _latex_to_plain(correct_answer)
                ap = container.add_paragraph()
                ap.paragraph_format.left_indent = Pt(24)
                ra = ap.add_run("Answer: ")
                ra.bold = True
                ra.font.size = Pt(10)
                ra.font.color.rgb = _rgb(tpl['correct'])
                rv = ap.add_run(correct)
                rv.font.size = Pt(10)
                rv.font.color.rgb = _rgb(tpl['correct'])

        if include_explanations:
            exp = _latex_to_plain(q.get('explanation', ''))
            if exp:
                ep = container.add_paragraph()
                ep.paragraph_format.left_indent = Pt(24)
                re2 = ep.add_run("Explanation: ")
                re2.bold = True
                re2.font.size = Pt(8)
                re2.font.color.rgb = _rgb(tpl['muted'])
                rv2 = ep.add_run(exp)
                rv2.font.size = Pt(8)
                rv2.font.color.rgb = _rgb(tpl['muted'])

    if has_sec:
        grouped = _group_by_section(questions)
        last_section_title = None

        for sec_key in sec_order:
            sec_qs = grouped.get(sec_key, [])
            if not sec_qs:
                continue

            sec_meta = sec_meta_dict.get(sec_key, {})
            current_title = sec_meta.get('title', sec_key)
            sec_letter = sec_key[:1]
            sec_color = _section_color(tpl, sec_letter)

            if current_title != last_section_title:
                doc.add_paragraph(rule_char * rule_len)
                h = doc.add_heading(current_title, level=1)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for hr in h.runs:
                    hr.font.color.rgb = _rgb(sec_color)
                    hr.font.name = tpl['docx_font']
                last_section_title = current_title

            sub_h = doc.add_paragraph()
            sub_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_r = sub_h.add_run(sec_meta.get('subtitle', ''))
            sub_r.font.size = Pt(9)
            sub_r.font.color.rgb = _rgb(tpl['muted'])
            sub_r.font.name = tpl['docx_font']

            inst_p = doc.add_paragraph()
            inst_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            inst_r = inst_p.add_run(sec_meta.get('instruction', ''))
            inst_r.font.size = Pt(9)
            inst_r.font.color.rgb = _rgb(tpl['muted'])
            inst_r.font.name = tpl['docx_font']
            inst_r.italic = True

            main_qs = [q for q in sec_qs if not q.get('_is_or', False)]
            or_qs = [q for q in sec_qs if q.get('_is_or', False)]
            or_queue = list(or_qs)

            for i, q in enumerate(main_qs):
                q_num += 1
                _render_q_docx(q, q_num, sec_letter)

                if or_queue:
                    or_q = or_queue.pop(0)
                    _add_docx_or_separator(doc)
                    _render_q_docx(or_q, q_num, sec_letter)

                if tpl.get('card_style') == 'flat' and (i < len(main_qs) - 1 or or_queue):
                    _add_docx_separator(doc)

        unsectioned = grouped.get('NONE', [])
        if unsectioned:
            doc.add_paragraph(rule_char * rule_len)
            h = doc.add_heading("Additional Questions", level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for hr in h.runs:
                hr.font.name = tpl['docx_font']

            sub_h = doc.add_paragraph()
            sub_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_r = sub_h.add_run("(Added by teacher)")
            sub_r.font.size = Pt(9)
            sub_r.font.color.rgb = _rgb(tpl['muted'])
            sub_r.font.name = tpl['docx_font']
            sub_r.italic = True

            for i, q in enumerate(unsectioned):
                q_num += 1
                _render_q_docx(q, q_num, None)
                if tpl.get('card_style') == 'flat' and i < len(unsectioned) - 1:
                    _add_docx_separator(doc)

    else:
        for i, q in enumerate(questions):
            q_num += 1
            _render_q_docx(q, q_num, None)
            if tpl.get('card_style') == 'flat' and i < len(questions) - 1:
                _add_docx_separator(doc)

    if include_answers and not include_explanations:
        doc.add_page_break()
        h = doc.add_heading("Answer Key", level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for hr in h.runs:
            hr.font.color.rgb = _rgb(tpl['primary'])
            hr.font.name = tpl['docx_font']

        q_num_ak = 0
        all_qs_ordered = []
        if has_sec:
            grouped = _group_by_section(questions)
            for sec_key in sec_order:
                all_qs_ordered.extend(grouped.get(sec_key, []))
            all_qs_ordered.extend(grouped.get('NONE', []))
        else:
            all_qs_ordered = questions

        for q in all_qs_ordered:
            q_num_ak += 1
            raw_table = q.get('answer_table') or q.get('answerTable')
            if raw_table and isinstance(raw_table, dict):
                p = doc.add_paragraph()
                p.add_run(f"Q{q_num_ak}. ").bold = True
                _render_answer_table_docx(doc, raw_table)
            else:
                correct = _latex_to_plain(q.get('correctAnswer', q.get('correct_answer', '')))
                p = doc.add_paragraph()
                p.add_run(f"Q{q_num_ak}. ").bold = True
                p.add_run(correct)

    doc.add_paragraph()
    ft = doc.add_paragraph()
    ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ft.add_run(f"Generated by a4ai · {board} {subject} Class {class_grade} · {display_date}")
    r.font.size = Pt(8)
    r.font.color.rgb = _rgb(tpl['light_muted'])
    r.font.name = tpl['docx_font']

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ═══════════════════════════════════════════════════════════════════════
# v14 — INSTITUTE-PAPER LAYOUT  (used by "colorful" template)
# ═══════════════════════════════════════════════════════════════════════
#
# Reference institute exam-paper format:
#   • Institute name header  →  "Class X — Subject"  →  optional Topic line
#   • Teacher / Max-Marks / Time / Date meta row
#   • Signature multi-color section rule
#   • Inline "[marks]" at end of each question (NOT a separate column)
#   • MCQ options rendered as (a)/(b) two-column pairs
#   • Plain "SECTION A (description)" headings
#   • "— All the Best —" footer
#
# Shared LaTeX/table/image helpers are reused; only the layout differs.


def _institute_section_heading(sec_key: str, meta_dict: dict) -> tuple:
    """Return (title, description) for an institute-style section heading."""
    meta = meta_dict.get(sec_key, {}) if meta_dict else {}
    title = meta.get("title", f"Section {sec_key}")
    subtitle = (meta.get("subtitle", "") or "").strip()
    # normalise "(1 mark each — MCQ / Assertion-Reason)" -> clean parenthetical
    desc = subtitle.strip("() ")
    return title.upper(), desc


def _institute_multicolor_rule_pdf(W: float, tpl: dict, thickness: float = 3.5):
    """A thin horizontal rule split into the template's section accent colors."""
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import Table, TableStyle

    sc = tpl.get("section_colors") or {}
    order = ["A", "B", "C", "D", "E", "F"]
    colors = [sc.get(k, tpl["primary"]) for k in order if sc.get(k)]
    if not colors:
        colors = [tpl["primary"]]

    n = len(colors)
    seg_w = W / n
    t = Table([[""] * n], colWidths=[seg_w] * n, rowHeights=[thickness])
    cmds = [
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
    ]
    for i, c in enumerate(colors):
        cmds.append(("BACKGROUND", (i, 0), (i, 0), HexColor(c)))
    t.setStyle(TableStyle(cmds))
    return t


def _generate_pdf_institute(
    questions, exam_title, board, class_grade, subject,
    include_answers, include_explanations, logo_base64, paper_date, tpl,
    teacher_name=None, institute_name=None, duration=None, topic=None,
) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, HRFlowable, Image as RLImage, KeepTogether,
    )

    buffer = io.BytesIO()
    top_m, bottom_m, left_m, right_m = tpl["margins_cm"]
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        topMargin=top_m * cm, bottomMargin=bottom_m * cm,
        leftMargin=left_m * cm, rightMargin=right_m * cm,
    )

    styles = getSampleStyleSheet()
    W = A4[0] - (left_m + right_m) * cm
    fb, fbd = tpl["font_body"], tpl["font_bold"]

    ist_styles = {
        "InstName":   dict(parent=styles["Title"], fontSize=18, leading=21, spaceAfter=1, alignment=TA_CENTER, textColor=HexColor(tpl["primary"]), fontName=fbd),
        "InstExam":   dict(parent=styles["Normal"], fontSize=12.5, leading=15, spaceAfter=1, alignment=TA_CENTER, textColor=HexColor(tpl["primary"]), fontName=fbd),
        "InstClass":  dict(parent=styles["Normal"], fontSize=11, leading=13, spaceAfter=1, alignment=TA_CENTER, textColor=HexColor(tpl["secondary"]), fontName=fb),
        "InstTopic":  dict(parent=styles["Normal"], fontSize=10, leading=12, spaceAfter=1, alignment=TA_CENTER, textColor=HexColor(tpl["muted"]), fontName=fb),
        "MetaL":      dict(parent=styles["Normal"], fontSize=10, leading=14, alignment=TA_LEFT, textColor=HexColor(tpl["secondary"]), fontName=fb),
        "MetaR":      dict(parent=styles["Normal"], fontSize=10, leading=14, alignment=TA_RIGHT, textColor=HexColor(tpl["secondary"]), fontName=fb),
        "GenInst":    dict(parent=styles["Normal"], fontSize=9, leading=12, alignment=TA_LEFT, textColor=HexColor(tpl["muted"]), fontName=fb),
        "SecHead":    dict(parent=styles["Heading2"], fontSize=11.5, leading=14, spaceBefore=12, spaceAfter=1, alignment=TA_LEFT, fontName=fbd),
        "QText":      dict(parent=styles["Normal"], fontSize=10.5, leading=13.5, spaceBefore=5, spaceAfter=2, alignment=TA_LEFT, textColor=HexColor("#1f1f3a"), fontName=fb),
        "Opt":        dict(parent=styles["Normal"], fontSize=10, leading=13, alignment=TA_LEFT, textColor=HexColor("#2b2b45"), fontName=fb),
        "OptCorrect": dict(parent=styles["Normal"], fontSize=10, leading=13, alignment=TA_LEFT, textColor=HexColor(tpl["correct"]), fontName=fbd),
        "Ans":        dict(parent=styles["Normal"], fontSize=9.5, leading=12, alignment=TA_LEFT, textColor=HexColor(tpl["correct"]), fontName=fbd),
        "Expl":       dict(parent=styles["Normal"], fontSize=9, leading=11.5, alignment=TA_LEFT, textColor=HexColor(tpl["muted"]), fontName=fb),
        "OrText":     dict(parent=styles["Normal"], fontSize=10, leading=13, alignment=TA_CENTER, textColor=HexColor(tpl["secondary"]), fontName=fbd, spaceBefore=3, spaceAfter=3),
        "Footer":     dict(parent=styles["Normal"], fontSize=8, alignment=TA_CENTER, textColor=HexColor(tpl["light_muted"]), fontName=fb),
        "AllBest":    dict(parent=styles["Normal"], fontSize=11, leading=14, alignment=TA_CENTER, textColor=HexColor(tpl["primary"]), fontName=fbd, spaceBefore=10),
    }
    # reportlab needs 'Option'/'AnswerLine' keys for shared table helpers
    for name, props in ist_styles.items():
        try:
            styles.add(ParagraphStyle(name=name, **props))
        except KeyError:
            pass
    # alias styles that shared table renderers look up by name
    for shared_name, src in (("Option", "Opt"), ("AnswerLine", "Ans"), ("QText", "QText")):
        if shared_name not in styles:
            try:
                styles.add(ParagraphStyle(name=shared_name, parent=styles[src]))
            except Exception:
                pass

    story = []
    display_date = _format_date_for_display(paper_date)
    total_marks = sum(q.get("marks", 1) for q in questions)

    # ── Header ──────────────────────────────────────────────────────
    if logo_base64:
        try:
            lb = logo_base64.split(",", 1)[1] if "," in logo_base64 else logo_base64
            logo_img = RLImage(io.BytesIO(base64.b64decode(lb)), width=1.6 * cm, height=1.6 * cm)
            logo_img.hAlign = "CENTER"
            story.append(logo_img)
            story.append(Spacer(1, 2))
        except Exception as e:
            logger.warning(f"Institute logo failed: {e}")

    header_name = (institute_name or "").strip() or (exam_title or "Test Paper")
    story.append(Paragraph(header_name, styles["InstName"]))

    # If both institute + exam title given and they differ, show exam title too
    if institute_name and exam_title and exam_title.strip() and exam_title.strip().lower() != header_name.strip().lower():
        story.append(Paragraph(exam_title.strip(), styles["InstExam"]))

    story.append(Paragraph(f"Class {class_grade} &nbsp;•&nbsp; {subject} &nbsp;•&nbsp; {board}", styles["InstClass"]))
    if topic and str(topic).strip():
        story.append(Paragraph(f"<i>Topic: {str(topic).strip()}</i>", styles["InstTopic"]))

    story.append(Spacer(1, 5))
    story.append(_institute_multicolor_rule_pdf(W, tpl))
    story.append(Spacer(1, 5))

    # ── Meta row (Teacher | Max Marks / Time / Date) ────────────────
    teacher_disp = (teacher_name or "").strip() or "______________"
    left_lines = [f"<b>Teacher:</b> {teacher_disp}"]
    right_lines = [f"<b>Max Marks:</b> {total_marks}"]
    if duration and str(duration).strip():
        right_lines.append(f"<b>Time:</b> {str(duration).strip()}")
    right_lines.append(f"<b>Date:</b> {display_date}")

    meta_left = Paragraph("<br/>".join(left_lines), styles["MetaL"])
    meta_right = Paragraph("<br/>".join(right_lines), styles["MetaR"])
    meta_tbl = Table([[meta_left, meta_right]], colWidths=[W * 0.55, W * 0.45])
    meta_tbl.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
    ]))
    story.append(meta_tbl)
    story.append(Spacer(1, 4))
    story.append(HRFlowable(width="100%", thickness=0.75, color=HexColor(tpl["border"]), spaceAfter=5))

    # ── Compact general instructions (single line, institute style) ─
    story.append(Paragraph(
        "<b>General Instructions:</b> All questions are compulsory. "
        "Marks for each question are indicated against it. "
        "Write answers neatly in the space provided.",
        styles["GenInst"],
    ))
    story.append(Spacer(1, 4))

    labels_lower = ["a", "b", "c", "d", "e", "f"]

    def _mcq_two_column(options, correct_answer):
        """Render MCQ options as (a)/(b) two-column pairs."""
        cells = []
        for idx, opt in enumerate(options):
            letter = labels_lower[idx] if idx < len(labels_lower) else str(idx + 1)
            opt_clean = re.sub(r'^[A-Fa-f][).\s]+\s*', '', _latex_to_paragraph(opt)).strip()
            is_correct = False
            if include_answers and correct_answer:
                ca = correct_answer.strip()
                if ca.upper().startswith(letter.upper()) or opt.strip() == ca.strip():
                    is_correct = True
            style = styles["OptCorrect"] if is_correct else styles["Opt"]
            cells.append(Paragraph(f"({letter}) {opt_clean}", style))

        # pack into 2-column rows
        rows = []
        for i in range(0, len(cells), 2):
            left = cells[i]
            right = cells[i + 1] if i + 1 < len(cells) else ""
            rows.append([left, right])
        if not rows:
            return []
        t = Table(rows, colWidths=[W * 0.5, W * 0.5])
        t.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 1),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
            ("LEFTPADDING", (0, 0), (-1, -1), 14),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ]))
        return [Spacer(1, 1), t]

    def _render_question_institute(q, q_num):
        elements = []
        raw_text = q.get("text", "")
        marks = q.get("marks", 1)
        marks_tag = f'<font color="{tpl["muted"]}"><b>[{marks}]</b></font>'

        question_table = _get_question_table(q)
        if question_table:
            raw_text = _strip_markdown_table_from_text(raw_text)

        segments = _split_text_and_tables(raw_text)
        first_text = ""
        for seg in segments:
            if seg["type"] == "text" and seg["content"]:
                first_text = _latex_to_paragraph(seg["content"])
                break
        if not first_text:
            first_text = _latex_to_paragraph(raw_text)

        # inline marks at end of question line (no separate column)
        elements.append(Paragraph(f"<b>{q_num}.</b> {first_text} &nbsp;{marks_tag}", styles["QText"]))

        if question_table:
            elements.extend(_render_question_table_pdf(question_table, styles, W))

        image_url = _get_image_url(q)
        if image_url:
            elements.extend(_render_manual_question_image_pdf(image_url, W))

        first_skipped = False
        for seg in segments:
            if seg["type"] == "text":
                if not first_skipped:
                    first_skipped = True
                    continue
                content = _latex_to_paragraph(seg["content"])
                if content:
                    elements.append(Paragraph(content, styles["QText"]))
            elif seg["type"] == "table":
                if question_table:
                    continue
                hdrs, rws = seg["content"]
                elements.extend(_render_inline_table_pdf(hdrs, rws, styles, W, tpl))

        options = q.get("options", [])
        correct_answer = q.get("correctAnswer", q.get("correct_answer", ""))

        if options:
            elements.extend(_mcq_two_column(options, correct_answer))
        else:
            fmt = q.get("format", "mcq")
            if not include_answers:
                gap = {"short_answer": 20, "long_answer": 44,
                       "journal_entry": 52, "ledger": 52, "trial_balance": 52,
                       "image": 22}.get(fmt, 0)
                if gap:
                    elements.append(Spacer(1, gap))

        if include_answers and include_explanations:
            raw_table = q.get("answer_table") or q.get("answerTable")
            if raw_table and isinstance(raw_table, dict):
                elements.extend(_render_answer_table_pdf(raw_table, styles, W))
            elif not options:
                ans = _latex_to_paragraph(correct_answer)
                elements.append(Paragraph(f"<b>Ans:</b> {ans}", styles["Ans"]))
        elif include_answers and not options:
            ans = _latex_to_paragraph(correct_answer)
            elements.append(Paragraph(f"<b>Ans:</b> {ans}", styles["Ans"]))

        if include_explanations:
            exp = _latex_to_paragraph(q.get("explanation", ""))
            if exp:
                elements.append(Paragraph(f"<b>Explanation:</b> {exp}", styles["Expl"]))

        return elements

    def _section_heading_flowables(sec_key, meta_dict):
        title, desc = _institute_section_heading(sec_key, meta_dict)
        accent = _section_color(tpl, sec_key[:1])
        head = f'<font color="{accent}">{title}</font>'
        if desc:
            head += f'  <font color="{tpl["muted"]}" size="9">({desc})</font>'
        return [
            Spacer(1, 6),
            Paragraph(head, styles["SecHead"]),
            HRFlowable(width="100%", thickness=1.2, color=HexColor(accent), spaceAfter=4),
        ]

    # ── Body ────────────────────────────────────────────────────────
    sec_order, sec_meta_dict = _get_section_order(questions)
    has_sec = sec_order is not None
    q_num = 0

    if has_sec:
        grouped = _group_by_section(questions)
        last_title = None
        for sec_key in sec_order:
            sec_qs = grouped.get(sec_key, [])
            if not sec_qs:
                continue
            meta = sec_meta_dict.get(sec_key, {})
            current_title = meta.get("title", sec_key)
            if current_title != last_title:
                story.extend(_section_heading_flowables(sec_key, sec_meta_dict))
                last_title = current_title

            main_qs = [q for q in sec_qs if not q.get("_is_or", False)]
            or_qs = [q for q in sec_qs if q.get("_is_or", False)]
            or_queue = list(or_qs)

            for q in main_qs:
                q_num += 1
                story.append(KeepTogether(_render_question_institute(q, q_num)))
                if or_queue:
                    or_q = or_queue.pop(0)
                    story.append(Paragraph("OR", styles["OrText"]))
                    story.append(KeepTogether(_render_question_institute(or_q, q_num)))

            for or_q in or_queue:
                q_num += 1
                story.append(Paragraph("OR", styles["OrText"]))
                story.append(KeepTogether(_render_question_institute(or_q, q_num)))

        unsectioned = grouped.get("NONE", [])
        if unsectioned:
            accent = _section_color(tpl, "F")
            story.append(Spacer(1, 6))
            story.append(Paragraph(f'<font color="{accent}">ADDITIONAL QUESTIONS</font>', styles["SecHead"]))
            story.append(HRFlowable(width="100%", thickness=1.2, color=HexColor(accent), spaceAfter=4))
            for q in unsectioned:
                q_num += 1
                story.append(KeepTogether(_render_question_institute(q, q_num)))
    else:
        for q in questions:
            q_num += 1
            story.append(KeepTogether(_render_question_institute(q, q_num)))

    # ── Answer key (answers-only mode) ──────────────────────────────
    if include_answers and not include_explanations:
        story.append(PageBreak())
        story.append(Paragraph("Answer Key", styles["InstName"]))
        story.append(Spacer(1, 4))
        story.append(_institute_multicolor_rule_pdf(W, tpl))
        story.append(Spacer(1, 6))

        all_qs = []
        if has_sec:
            grouped = _group_by_section(questions)
            for sec_key in sec_order:
                all_qs.extend(grouped.get(sec_key, []))
            all_qs.extend(grouped.get("NONE", []))
        else:
            all_qs = questions

        for i, q in enumerate(all_qs, 1):
            raw_table = q.get("answer_table") or q.get("answerTable")
            if raw_table and isinstance(raw_table, dict):
                story.append(Paragraph(f"<b>{i}.</b>", styles["QText"]))
                story.extend(_render_answer_table_pdf(raw_table, styles, W))
            else:
                correct = _latex_to_paragraph(q.get("correctAnswer", q.get("correct_answer", "")))
                story.append(Paragraph(f"<b>{i}.</b> {correct}", styles["QText"]))

    # ── Footer ──────────────────────────────────────────────────────
    story.append(Paragraph("— All the Best —", styles["AllBest"]))
    story.append(Spacer(1, 6))
    story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor(tpl["border"]), spaceAfter=4))
    story.append(Paragraph(f"Generated by a4ai · {board} {subject} Class {class_grade} · {display_date}", styles["Footer"]))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def _institute_multicolor_rule_docx(doc, tpl: dict):
    """A thin horizontal rule split into the template's section accent colors (DOCX)."""
    from docx.shared import Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
    from docx.oxml.ns import qn

    sc = tpl.get("section_colors") or {}
    order = ["A", "B", "C", "D", "E", "F"]
    colors = [sc.get(k) for k in order if sc.get(k)]
    if not colors:
        colors = [tpl["primary"]]

    n = len(colors)
    table = doc.add_table(rows=1, cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    row = table.rows[0]
    row.height = Pt(4)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    for i, c in enumerate(colors):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(" ")
        run.font.size = Pt(1)
        tcPr = cell._element.get_or_add_tcPr()
        shd = tcPr.makeelement(qn('w:shd'), {qn('w:fill'): _hexnc(c), qn('w:val'): 'clear'})
        tcPr.append(shd)
        # zero cell margins
        tcMar = tcPr.makeelement(qn('w:tcMar'), {})
        for side in ('top', 'bottom', 'start', 'end'):
            m = tcMar.makeelement(qn(f'w:{side}'), {qn('w:w'): '0', qn('w:type'): 'dxa'})
            tcMar.append(m)
        tcPr.append(tcMar)


def _generate_docx_institute(
    questions, exam_title, board, class_grade, subject,
    include_answers, include_explanations, logo_base64, paper_date, tpl,
    teacher_name=None, institute_name=None, duration=None, topic=None,
) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()
    font_name = tpl["docx_font"]

    try:
        normal_style = doc.styles["Normal"]
        normal_style.font.name = font_name
        rpr = normal_style.element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = rpr.makeelement(qn('w:rFonts'), {})
            rpr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)
    except Exception as e:
        logger.warning(f"Institute docx font set failed: {e}")

    top_m, bottom_m, left_m, right_m = tpl["margins_cm"]
    for section in doc.sections:
        section.top_margin = Cm(top_m)
        section.bottom_margin = Cm(bottom_m)
        section.left_margin = Cm(left_m)
        section.right_margin = Cm(right_m)

    display_date = _format_date_for_display(paper_date)
    total_marks = sum(q.get("marks", 1) for q in questions)
    labels_lower = ["a", "b", "c", "d", "e", "f"]

    def _center_run(text, size, color_hex, bold=True, italic=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.color.rgb = _rgb(color_hex)
        r.font.name = font_name
        return p

    # ── Header ──────────────────────────────────────────────────────
    if logo_base64:
        try:
            lb = logo_base64.split(",", 1)[1] if "," in logo_base64 else logo_base64
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(io.BytesIO(base64.b64decode(lb)), width=Cm(1.8))
        except Exception as e:
            logger.warning(f"Institute docx logo failed: {e}")

    header_name = (institute_name or "").strip() or (exam_title or "Test Paper")
    _center_run(header_name, 18, tpl["primary"], bold=True)

    if institute_name and exam_title and exam_title.strip() and exam_title.strip().lower() != header_name.strip().lower():
        _center_run(exam_title.strip(), 12.5, tpl["primary"], bold=True)

    _center_run(f"Class {class_grade}  •  {subject}  •  {board}", 11, tpl["secondary"], bold=False)
    if topic and str(topic).strip():
        _center_run(f"Topic: {str(topic).strip()}", 10, tpl["muted"], bold=False, italic=True)

    _institute_multicolor_rule_docx(doc, tpl)

    # ── Meta row ────────────────────────────────────────────────────
    teacher_disp = (teacher_name or "").strip() or "______________"
    meta_tbl = doc.add_table(rows=1, cols=2)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    lc = meta_tbl.rows[0].cells[0]
    lc.text = ""
    lp = lc.paragraphs[0]
    lr = lp.add_run("Teacher: ")
    lr.bold = True; lr.font.size = Pt(10); lr.font.name = font_name
    lr2 = lp.add_run(teacher_disp)
    lr2.font.size = Pt(10); lr2.font.name = font_name

    rc = meta_tbl.rows[0].cells[1]
    rc.text = ""
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_bits = [("Max Marks: ", str(total_marks))]
    if duration and str(duration).strip():
        right_bits.append(("Time: ", str(duration).strip()))
    right_bits.append(("Date: ", display_date))
    for k, (lbl, val) in enumerate(right_bits):
        rb = rp.add_run(lbl); rb.bold = True; rb.font.size = Pt(10); rb.font.name = font_name
        rv = rp.add_run(val); rv.font.size = Pt(10); rv.font.name = font_name
        if k < len(right_bits) - 1:
            sep = rp.add_run("    "); sep.font.size = Pt(10)

    # thin separator line
    sep_p = doc.add_paragraph()
    sep_r = sep_p.add_run("─" * 60)
    sep_r.font.size = Pt(7)
    sep_r.font.color.rgb = _rgb(tpl["border"])

    # ── General instructions (single line) ──────────────────────────
    gi = doc.add_paragraph()
    gir = gi.add_run("General Instructions: ")
    gir.bold = True; gir.font.size = Pt(9); gir.font.name = font_name
    gir.font.color.rgb = _rgb(tpl["muted"])
    giv = gi.add_run("All questions are compulsory. Marks for each question are indicated against it. "
                     "Write answers neatly in the space provided.")
    giv.font.size = Pt(9); giv.font.name = font_name
    giv.font.color.rgb = _rgb(tpl["muted"])

    def _mcq_two_column_docx(options, correct_answer):
        n = len(options)
        if n == 0:
            return
        nrows = (n + 1) // 2
        table = doc.add_table(rows=nrows, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for idx, opt in enumerate(options):
            letter = labels_lower[idx] if idx < len(labels_lower) else str(idx + 1)
            opt_clean = re.sub(r'^[A-Fa-f][).\s]+\s*', '', _latex_to_plain(opt)).strip()
            is_correct = bool(include_answers and correct_answer
                              and correct_answer.strip().upper().startswith(letter.upper()))
            row_i, col_i = idx // 2, idx % 2
            cell = table.rows[row_i].cells[col_i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.left_indent = Pt(14)
            r = p.add_run(f"({letter}) {opt_clean}")
            r.font.size = Pt(10)
            r.font.name = font_name
            if is_correct:
                r.bold = True
                r.font.color.rgb = _rgb(tpl["correct"])

    def _render_q_docx_institute(q, q_num):
        raw_text = q.get("text", "")
        marks = q.get("marks", 1)

        question_table = _get_question_table(q)
        if question_table:
            raw_text = _strip_markdown_table_from_text(raw_text)

        segments = _split_text_and_tables(raw_text)
        first_text = ""
        for seg in segments:
            if seg["type"] == "text" and seg["content"]:
                first_text = _latex_to_plain(seg["content"])
                break
        if not first_text:
            first_text = _latex_to_plain(raw_text)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(2)
        rq = p.add_run(f"{q_num}. ")
        rq.bold = True; rq.font.size = Pt(10.5); rq.font.name = font_name
        rt = p.add_run(first_text)
        rt.font.size = Pt(10.5); rt.font.name = font_name
        rm = p.add_run(f"   [{marks}]")
        rm.bold = True; rm.font.size = Pt(9); rm.font.name = font_name
        rm.font.color.rgb = _rgb(tpl["muted"])

        if question_table:
            _render_question_table_docx(doc, question_table)

        image_url = _get_image_url(q)
        if image_url:
            _render_manual_question_image_docx(doc, image_url)

        first_skipped = False
        for seg in segments:
            if seg["type"] == "text":
                if not first_skipped:
                    first_skipped = True
                    continue
                content = _latex_to_plain(seg["content"])
                if content:
                    cp = doc.add_paragraph()
                    cr = cp.add_run(content)
                    cr.font.size = Pt(10.5); cr.font.name = font_name
            elif seg["type"] == "table":
                if question_table:
                    continue
                hdrs, rws = seg["content"]
                _render_inline_table_docx(doc, hdrs, rws)

        options = q.get("options", [])
        correct_answer = q.get("correctAnswer", q.get("correct_answer", ""))

        if options:
            _mcq_two_column_docx(options, correct_answer)
        else:
            fmt = q.get("format", "mcq")
            if not include_answers:
                blanks = {"short_answer": 2, "long_answer": 4,
                          "journal_entry": 5, "ledger": 5, "trial_balance": 5,
                          "image": 2}.get(fmt, 0)
                for _ in range(blanks):
                    doc.add_paragraph()

        if include_answers and include_explanations:
            raw_table = q.get("answer_table") or q.get("answerTable")
            if raw_table and isinstance(raw_table, dict):
                _render_answer_table_docx(doc, raw_table)
            elif not options:
                ap = doc.add_paragraph()
                ar = ap.add_run("Ans: "); ar.bold = True
                ar.font.size = Pt(10); ar.font.color.rgb = _rgb(tpl["correct"]); ar.font.name = font_name
                av = ap.add_run(_latex_to_plain(correct_answer))
                av.font.size = Pt(10); av.font.color.rgb = _rgb(tpl["correct"]); av.font.name = font_name
        elif include_answers and not options:
            ap = doc.add_paragraph()
            ar = ap.add_run("Ans: "); ar.bold = True
            ar.font.size = Pt(10); ar.font.color.rgb = _rgb(tpl["correct"]); ar.font.name = font_name
            av = ap.add_run(_latex_to_plain(correct_answer))
            av.font.size = Pt(10); av.font.color.rgb = _rgb(tpl["correct"]); av.font.name = font_name

        if include_explanations:
            exp = _latex_to_plain(q.get("explanation", ""))
            if exp:
                ep = doc.add_paragraph()
                er = ep.add_run("Explanation: "); er.bold = True
                er.font.size = Pt(8.5); er.font.color.rgb = _rgb(tpl["muted"]); er.font.name = font_name
                ev = ep.add_run(exp)
                ev.font.size = Pt(8.5); ev.font.color.rgb = _rgb(tpl["muted"]); ev.font.name = font_name

    def _section_heading_docx(sec_key, meta_dict):
        title, desc = _institute_section_heading(sec_key, meta_dict)
        accent = _section_color(tpl, sec_key[:1])
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(title)
        r.bold = True; r.font.size = Pt(11.5); r.font.name = font_name
        r.font.color.rgb = _rgb(accent)
        if desc:
            rd = p.add_run(f"  ({desc})")
            rd.font.size = Pt(9); rd.font.name = font_name
            rd.font.color.rgb = _rgb(tpl["muted"])
        line = doc.add_paragraph()
        lr = line.add_run("─" * 60)
        lr.font.size = Pt(7); lr.font.color.rgb = _rgb(accent)

    def _or_docx():
        op = doc.add_paragraph()
        op.alignment = WD_ALIGN_PARAGRAPH.CENTER
        orr = op.add_run("OR")
        orr.bold = True; orr.font.size = Pt(10); orr.font.name = font_name
        orr.font.color.rgb = _rgb(tpl["secondary"])

    # ── Body ────────────────────────────────────────────────────────
    sec_order, sec_meta_dict = _get_section_order(questions)
    has_sec = sec_order is not None
    q_num = 0

    if has_sec:
        grouped = _group_by_section(questions)
        last_title = None
        for sec_key in sec_order:
            sec_qs = grouped.get(sec_key, [])
            if not sec_qs:
                continue
            meta = sec_meta_dict.get(sec_key, {})
            current_title = meta.get("title", sec_key)
            if current_title != last_title:
                _section_heading_docx(sec_key, sec_meta_dict)
                last_title = current_title

            main_qs = [q for q in sec_qs if not q.get("_is_or", False)]
            or_qs = [q for q in sec_qs if q.get("_is_or", False)]
            or_queue = list(or_qs)

            for q in main_qs:
                q_num += 1
                _render_q_docx_institute(q, q_num)
                if or_queue:
                    or_q = or_queue.pop(0)
                    _or_docx()
                    _render_q_docx_institute(or_q, q_num)

            for or_q in or_queue:
                q_num += 1
                _or_docx()
                _render_q_docx_institute(or_q, q_num)

        unsectioned = grouped.get("NONE", [])
        if unsectioned:
            accent = _section_color(tpl, "F")
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            r = p.add_run("ADDITIONAL QUESTIONS")
            r.bold = True; r.font.size = Pt(11.5); r.font.name = font_name
            r.font.color.rgb = _rgb(accent)
            line = doc.add_paragraph()
            lr = line.add_run("─" * 60); lr.font.size = Pt(7); lr.font.color.rgb = _rgb(accent)
            for q in unsectioned:
                q_num += 1
                _render_q_docx_institute(q, q_num)
    else:
        for q in questions:
            q_num += 1
            _render_q_docx_institute(q, q_num)

    # ── Answer key (answers-only mode) ──────────────────────────────
    if include_answers and not include_explanations:
        doc.add_page_break()
        _center_run("Answer Key", 18, tpl["primary"], bold=True)
        _institute_multicolor_rule_docx(doc, tpl)

        all_qs = []
        if has_sec:
            grouped = _group_by_section(questions)
            for sec_key in sec_order:
                all_qs.extend(grouped.get(sec_key, []))
            all_qs.extend(grouped.get("NONE", []))
        else:
            all_qs = questions

        for i, q in enumerate(all_qs, 1):
            raw_table = q.get("answer_table") or q.get("answerTable")
            if raw_table and isinstance(raw_table, dict):
                p = doc.add_paragraph()
                p.add_run(f"{i}. ").bold = True
                _render_answer_table_docx(doc, raw_table)
            else:
                correct = _latex_to_plain(q.get("correctAnswer", q.get("correct_answer", "")))
                p = doc.add_paragraph()
                p.add_run(f"{i}. ").bold = True
                p.add_run(correct)

    # ── Footer ──────────────────────────────────────────────────────
    _center_run("— All the Best —", 11, tpl["primary"], bold=True)
    ft = doc.add_paragraph()
    ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ft.add_run(f"Generated by a4ai · {board} {subject} Class {class_grade} · {display_date}")
    r.font.size = Pt(8); r.font.color.rgb = _rgb(tpl["light_muted"]); r.font.name = font_name

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()